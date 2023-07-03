import pandas as pd
from docx import Document
import matplotlib.pyplot as plt
from docx.shared import Inches, Pt
import numpy as np
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io

H="threephase.csv"
tf=pd.read_csv("threephase.csv")
I="threephasevalue.csv"
tf2=pd.read_csv("threephasevalue.csv")


def threephase_result(tf, tf2):
    tf["Rated Line Voltage (V)"] = tf2["Rated Line Voltage (V)"]
    tf["Average Line Voltage (V)"] = round(
        (tf2["Voltage-L1L2 (V)"] + tf2["Voltage-L2L3 (V)"] + tf2["Voltage-L3L1 (V)"]) / 3, 2
    )

    tf["Average Phase Voltage (V)"] = (
        tf2["Voltage-L1N (V)"] + tf2["Voltage-L2N (V)"] + tf2["Voltage-L3N (V)"]
    ) / 3
    tf["Voltage Unbalance %"] = round(
        (
            (tf2["Voltage-L1N (V)"] - tf["Average Line Voltage (V)"])
            .abs()
            .where(
                (tf2["Voltage-L1N (V)"] - tf["Average Line Voltage (V)"]) > 0,
                (tf["Average Line Voltage (V)"] - tf2["Voltage-L1N (V)"]).abs(),
            )
            .max(axis=0)
            / tf["Average Line Voltage (V)"]
        )
        * 100,
        2,
    )
    tf["Voltage Result"] = np.where(tf["Voltage Unbalance %"] <= 10, "PASS", "FAIL")
    tf["Rated Phase Current (A)"] = tf2["Rated Phase Current (A)"]
    tf["Average Phase Current (A)"] = round(
        (tf2["Current-L1 (A)"] + tf2["Current-L2 (A)"] + tf2["Current-L3 (A)"]) / 3, 2
    )
    tf["Current Unbalance %"] = round(
        (
            (tf2["Current-L1 (A)"] - tf["Average Phase Current (A)"])
            .abs()
            .where(
                (tf2["Current-L1 (A)"] - tf["Average Phase Current (A)"]) > 0,
                (tf["Average Phase Current (A)"] - tf2["Current-L1 (A)"]).abs(),
            )
            .max(axis=0)
            / tf["Average Line Voltage (V)"]
        )
        * 100,
        2,
    )
    tf["Current Result"] = np.where(tf["Current Unbalance %"] <= 10, "PASS", "FAIL")
    tf["Voltage-NE (V)"] = tf2["Voltage-NE (V)"]
    tf["NEV Result"] = np.where(tf["Voltage-NE (V)"] <= 2, "PASS", "FAIL")
    tf["Zero Sum Current (mA)"] = tf2["Zero Sum Current (mA)"]
    tf["ZeroSum Result"] = np.where(
        tf["Zero Sum Current (mA)"] <= (tf["Rated Phase Current (A)"] / 5000 * 1000),
        "PASS",
        "FAIL",
    )
    return tf



def threephase_table(tf, doc):
    table_data = tf.values
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    table.autofit = False
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.5,
        3: 0.54,
        4: 0.42,
        5: 0.56,
        6: 0.5,
        7: 0.55,
        8: 0.46,
        9: 0.43,
        10: 0.52,
        11: 0.56,
        12: 0.46,
        13: 0.5,
        14: 0.46,
        15: 0.48,
        16: 0.5,
    }
    for j, col in enumerate(tf.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    for i, row in enumerate(table_data, start=1):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)
        font_size = 6
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc

def threephase_combined_graph(tf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = tf["Facility Area"]
    y = tf["Zero Sum Current (mA)"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Facility Area")
    plt.ylabel("Zero Sum Current (mA)")
    plt.title("Facility Area VS  Zero Sum Current (mA)")

    # Pie chart
    plt.subplot(122)
    plt.figure(figsize=(6, 6))
    result_counts = tf["ZeroSum Result"].value_counts()
    labels = result_counts.index
    values = result_counts.values
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
    plt.title("Three Phase Symmetry Test Results")
    plt.axis('equal')  # Equal aspect ratio ensures that the pie is drawn as a circle

    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined


def main():
    tf = pd.read_csv("threephase.csv")
    tf2 = pd.read_csv("threephasevalue.csv")
    tf = threephase_result(tf, tf2)
    tf.to_csv("main.csv", index=False)

    doc = Document()
    doc.add_heading("Three Phase Symmetry Test", 0)
    for section in doc.sections:
        section.left_margin = Inches(0.2)
    doc = threephase_table(tf, doc)
    graph_combined = threephase_combined_graph(tf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))
    doc.save("threephaseSymmetry.docx")

main()
