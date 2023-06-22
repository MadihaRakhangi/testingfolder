import pytest
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from residual import Resi_create_table, Resi_graph, Resi_result

test_rf = pd.read_csv("residual.csv")

@pytest.fixture
def sample_data():
    rf = {
        "Type": ["AC", "A"],
        "Test Current (mA)": [10, 20],
        "Rated Residual Operating Current,I?n (mA)": [20, 40],
        "Device Tripped": ["Yes", "Yes"],
        "Trip Time (ms)": ["100", "200"],
        "Result": ["Pass", "Pass"]
    }
    return pd.DataFrame(rf)

def test_Resi_result():
    assert Resi_result("AC", 1, 1, "No", "-") == "Fail"
    assert Resi_result("AC", 0.5, 1, "Yes", 100) == "Fail"
    assert Resi_result("A", 1, 1, "Yes", 150) == "Pass"
    assert Resi_result("AC", 1, 1, "No", "Invalid") == "Fail"


def test_create_table():
    rf = pd.DataFrame({                               #create  a test dataframe
        "Type": ["AC", "A"],
        "Test Current (mA)": [10, 20],
        "Rated Residual Operating Current,I?n (mA)": [20, 40],
        "Device Tripped": ["Yes", "Yes"],
        "Trip Time (ms)": ["100", "200"],
        "Result": ["Pass", "Pass"]
    })

    doc = Document()                                  # Create a Document object
    doc =Resi_create_table(rf, doc)                      
    temp_file = "temp_doc.docx"                        # Save the document as a temporary file
    temp_file = "temp_doc.docx"
    doc.save(temp_file)
    doc_read = Document(temp_file)                     # Read the temporary file and check if the table exists
    tables = doc_read.tables
    assert len(tables) == 1                             # Check if the first table in the document matches the DataFrame shape
    table = tables[0]
    assert len(table.rows) == rf.shape[0] + 1            # Include header row
    assert len(table.columns) == rf.shape[1]          #include the result coloumn
    os.remove(temp_file) 

    
def test_Resi_graph(sample_data):
    graph = Resi_graph(sample_data)

    assert graph is not None