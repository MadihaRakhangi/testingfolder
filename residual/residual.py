import pandas as pd
import matplotlib.pyplot as plt
import docx
from docx import Document
import csv
import io
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
import numpy as np
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

E="residual.csv"
rf = pd.read_csv("residual.csv")



def residual_result(Type, Test_Current, Rated_OpCurrent, D_Tripped, Trip_Time):
    if Type == "AC":
        if Test_Current == 0.5 * Rated_OpCurrent:
            if D_Tripped == "No":
                return "Pass"
            else:
                return "Fail"
        elif Test_Current == 1 * Rated_OpCurrent and D_Tripped == "Yes":
            if Trip_Time <= 300:
                return "Pass"
            else:
                return "Fail"
        elif Test_Current == 2 * Rated_OpCurrent and D_Tripped == "Yes":
            if Trip_Time <= 150:
                return "Pass"
            else:
                return "Fail"
        elif Test_Current == 5 * Rated_OpCurrent and D_Tripped == "Yes":
            if Trip_Time <= 40:
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"
    elif Type == "A":
        if Test_Current == 0.5 * Rated_OpCurrent:
            if D_Tripped == "No":
                return "Pass"
            else:
                return "Fail"
        elif Test_Current == 1 * Rated_OpCurrent and D_Tripped == "Yes":
            if 130 <= Trip_Time <= 500:
                return "Pass"
            else:
                return "Fail"
        elif Test_Current == 2 * Rated_OpCurrent and D_Tripped == "Yes":
            if 60 <= Trip_Time <= 200:
                return "Pass"
            else:
                return "Fail"
        elif Test_Current == 5 * Rated_OpCurrent and D_Tripped == "Yes":
            if 50 <= Trip_Time <= 150:
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"
    else:
        return "Pass"


def residual_rang(length):
    res7 = []
    for row in range(0, length):
        Type= rf.iloc[row, 4]
        Test_Current = rf.iloc[row, 11]
        Rated_OpCurrent= rf.iloc[row, 8]
        D_Tripped = rf.iloc[row, 14] 
        Trip_Time = rf.iloc[row, 13]
        res7.append(residual_result(Type, Test_Current, Rated_OpCurrent, D_Tripped, Trip_Time))
    return res7


from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def Resi_create_table(rf, doc):
    rf["Result"] = np.NaN
    for index, row in rf.iterrows():
        trip_time_str = row["Trip Time (ms)"]
        if trip_time_str.isnumeric():
            trip_time = int(trip_time_str)
            result_val = residual_result(
                row["Type"],
                row["Test Current (mA)"],
                row["Rated Residual Operating Current,I?n (mA)"],
                row["Device Tripped"],
                trip_time,
            )
            rf.loc[index, "Result"] = result_val
        elif trip_time_str == "-":
            rf.loc[index, "Result"] = "Pass"
        else:
            rf.loc[index, "Result"] = "invalid"

    table_data = rf.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    table.autofit = False
    column_widths = {
        0: 0.25,
        1: 0.54,
        2: 0.55,
        3: 0.59,
        4: 0.48,
        5: 0.56,
        6: 0.54,
        7: 0.59,
        8: 0.37,
        9: 0.55,
        10: 0.42,
        11: 0.56,
        12: 0.48,
        13: 0.4,
        14: 0.49,
        15: 0.5,
        num_cols: 0.4,  # Width for Result column
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)
    
    # Add shading to the Result column based on the result value
    for i in range(1, num_rows + 1):
        result = table.cell(i, num_cols - 1).text
        cell = table.cell(i, num_cols - 1)
        if result == "Pass":
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="#5ac85a"/>'.format(nsdecls("w"))
            )  # Green color
            cell._tc.get_or_add_tcPr().append(shading_elm)
        else:
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="#dc0000"/>'.format(nsdecls("w"))
            )  # Red color
            cell._tc.get_or_add_tcPr().append(shading_elm)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc


def residual_combined_graph(rf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    y= rf["Trip curve type"]
    x = rf["No. of Poles"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.ylabel("Trip curve type")
    plt.xlabel("No. of Poles")
    plt.title("Residual Current Device Test Results")

    # Pie chart
    plt.subplot(122)
    # rf['Result'] = residual_rang(rf.shape[0])  # Ensure you have the residual_rang() function defined correctly
    result_counts = rf["Result"].value_counts()
    labels = result_counts.index
    values = result_counts.values
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", shadow=False, startangle=90, colors=colors)
    plt.title("Residual Test Results")
    plt.axis("equal")  # Equal aspect ratio ensures that the pie is drawn as a circle
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()


    

    return graph_combined




  
def main():
    rf = pd.read_csv("residual.csv")
    doc = Document()
    doc.add_heading("Residual Current Device Test", 0)
    for section in doc.sections:
        section.left_margin = Inches(0.2)
    doc = Resi_create_table(rf, doc)
    graph_combined = residual_combined_graph(rf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))
    doc.save("Residual.docx")


main()