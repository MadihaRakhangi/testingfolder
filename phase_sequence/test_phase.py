import pytest
from phase_seq import phase_graph,phase_pie,phase_table,phaserang,phase_result
import pandas as pd
import io
import os
import matplotlib as plt
from docx import Document


test_df=pd.read_csv("phasesequence.csv")

@pytest.fixture
def sample_dataframe():
                                                           
    df = pd.DataFrame({                                           # Define a sample dataframe for testing
        "Phase Sequence": ["RYB", "RBY"],
        "V-L2-N": ["238", "241"]
        
    })
    return df

def test_result():
    assert phase_result("RYB") == "CLOCKWISE"
    assert phase_result("RBY") == "ANTICLOCKWISE"

# def test_rang():
#     df = pd.DataFrame({                                  #create  a test dataframe
#          "Phase Sequence":["RYB","RBY"]
#     })
#     assert rang(df.shape[0]) == ["CLOCKWISE", "ANTICLOCKWISE"]

def test_rang():
    df = pd.DataFrame({
        "Phase Sequence": ["RYB", "RBY"]
    })
    assert phaserang(df) == ["CLOCKWISE", "ANTICLOCKWISE"]

def test_create_table():
    df = pd.DataFrame({                               #create  a test dataframe
        "Serial number": [1, 2],
        "Phase Sequence":["RYB","RBY"]
    })

    doc = Document()                                  # Create a Document object
    doc = phase_table(df, doc)                      
    temp_file = "temp_doc.docx"                        # Save the document as a temporary file
    temp_file = "temp_doc.docx"
    doc.save(temp_file)
    doc_read = Document(temp_file)                     # Read the temporary file and check if the table exists
    tables = doc_read.tables
    assert len(tables) == 1                             # Check if the first table in the document matches the DataFrame shape
    table = tables[0]
    assert len(table.rows) == df.shape[0] + 1            # Include header row
    assert len(table.columns) == df.shape[1]+1           #include the result coloumn
    os.remove(temp_file)    

def test_phase_graph(sample_dataframe):
    graph4 = phase_graph(sample_dataframe)
    # Add assertions to check the generated graph

def test_phase_pie(sample_dataframe):
    graph5 = phase_pie(sample_dataframe)
    # Add assertions to check the generated graph