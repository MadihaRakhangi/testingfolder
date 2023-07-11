import pandas as pd
import matplotlib.pyplot as plt
import docx
from docx import Document
import csv
import io
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

C="phasesequence.csv"
pf = pd.read_csv("phasesequence.csv")


def phase_result(phase_seq):
    if phase_seq == 'RYB':
        return "CLOCKWISE"
    else:
        return "ANTICLOCKWISE"

def phase_rang(pf):
    res3 = []
    phase_seqs = pf["Phase Sequence"]
    for seq in phase_seqs:
        if seq == "RYB":
            res3.append("CLOCKWISE")
        elif seq == "RBY":
            res3.append("ANTICLOCKWISE")
        else:
            res3.append("UNKNOWN")
    return res3

def phase_table(pf, doc):
    table_data = pf.iloc[:, :]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.2,
        1: 0.4,
        2: 0.4,
        3: 0.4,
        4: 0.6,
        5: 0.3,
        6: 0.3,
        7: 0.4,
        8: 0.4,
        9: 0.4,
        10: 0.4,
        11: 0.4,
        12: 0.6,
        num_cols: 0.8,  # Width for the "Result" column
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    for i, row in enumerate(table_data.itertuples(), start=0):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i + 1, j).text = str(value)
    results = phase_rang(pf)

    table.cell(0, num_cols).text = "Result"
    for i, result in enumerate(results, start=1):
        cell = table.cell(i, num_cols)
        cell.text = result
        
        if result == "CLOCKWISE":
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="#5ac85a"/>'.format(nsdecls("w"))
            )  # Green color
            cell._tc.get_or_add_tcPr().append(shading_elm)
        elif result == "ANTICLOCKWISE":
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="#dc0000"/>'.format(nsdecls("w"))
            )  # Red color
            cell._tc.get_or_add_tcPr().append(shading_elm)

        cell.width = Inches(column_widths[num_cols])  # Set width for the "Result" column

    font_size = 8
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc


def phase_combined_graph(pf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = pf["Phase Sequence"]
    y = pf["V-L3-N"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Phase Sequence")
    plt.ylabel("V-L3-N")
    plt.title("Phase Sequence by V-L3-N")

    # Pie chart
    plt.subplot(122)
    pf['Result'] = phase_rang(pf)  # Ensure you have the phase_rang() function defined correctly
    pf_counts = pf['Result'].value_counts()
    labels = pf_counts.index.tolist()
    values = pf_counts.values.tolist()
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
    plt.axis('equal')
    plt.title('Test Results')
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined



def main():
    pf = pd.read_csv("phasesequence.csv")
    doc = Document()
    doc.add_heading('Phase Sequence test', 0)
    doc = phase_table(pf, doc)
    graph_combined = phase_combined_graph(pf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))  
    doc.save("phasesequence.docx")

main()
