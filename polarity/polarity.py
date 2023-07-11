import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import io
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

D="pol.csv"
af = pd.read_csv("pol.csv")





def polarity_result(line_neutral):
    if line_neutral == 230:
        return "OK"
    else :
        return "REVERSE"
    

def polarityrang(length):
    res4 = []
    for row in range(0, length):
        line_neutral = af.iloc[row, 5]
        res4.append(polarity_result(line_neutral))
    return res4

def polarity_table(af, doc):
    table_data = af.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.55,
        3: 0.54,
        4: 0.38,
        5: 0.56,
        6: 0.5,
        7: 0.48,
        8: 0.71,
        num_cols: 0.8,  # Width for the "Result" column
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)
    Results = polarityrang(num_rows)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(column_widths[num_cols])  # Set width for the "Result" column
    for i in range(num_rows):
        res_index = i
        result = Results[res_index]
        cell = table.cell(i + 1, num_cols)
        cell.text = result
        if result == "OK":
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="#5ac85a"/>'.format(nsdecls("w"))
            )  # Green color
            cell._tc.get_or_add_tcPr().append(shading_elm)
        else:
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="#dc0000"/>'.format(nsdecls("w"))
            )  # Red color
            cell._tc.get_or_add_tcPr().append(shading_elm)

    font_size = 7

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc


def polarity_combined_graph(af):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = af["Type of Supply"]
    y = af["Line to Neutral Voltage (V)"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Type of Supply")
    plt.ylabel("Line to Neutral Voltage (V)")
    plt.title("Type of Supply VS Line to Neutral Voltage (V)")

    # Pie chart
    plt.subplot(122)
    af['Result'] = polarityrang(af.shape[0])  # Ensure you have the polarityrang() function defined correctly
    af_counts = af['Result'].value_counts()
    labels = af_counts.index.tolist()
    values = af_counts.values.tolist()
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
    plt.axis('equal')
    plt.title('Polarity Results')
    # Save the combined graph as bytes
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined





def main():
    doc = Document()
    doc.add_heading("POLARITY TEST", 0)
    doc = polarity_table(af, doc)
    graph_combined = polarity_combined_graph(af)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))
    doc.save("polarity.docx")


main()