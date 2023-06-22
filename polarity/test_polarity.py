import pytest
from polarity import polarity_graph , polarity_result , polarityrang , polarity_pie , polarity_table
import pandas as pd
import io
import os
import matplotlib as plt
from docx import Document

test_df=pd.read_csv("pol.csv")


@pytest.fixture
def sample_dataframe():
                                                           
    df = pd.DataFrame({                                           # Define a sample dataframe for testing
        "Line to Neutral Voltage (V)": [230, 220, 245],
        "Type of Supply":["AC","DC","AC"]
    })
    return df


def test_result():
    assert polarity_result(230) == "OK"
    assert polarity_result(245) == "REVERSE"

def test_rang():
    df = pd.DataFrame({                                  #create  a test dataframe
         "Line to Neutral Voltage (V)": [230, 220, 245]
    })
    assert polarityrang(df.shape[0]) == ["OK", "REVERSE", "REVERSE"]

def test_create_table():
    df = pd.DataFrame({                               #create  a test dataframe
        "Serial number": [1, 2, 3],
        "Line to Neutral Voltage (V)": [230, 220, 245],
        "Type of Supply":["AC","DC","AC"]
    })

    doc = Document()                                  # Create a Document object
    doc = polarity_table(df, doc)                      
    temp_file = "temp_doc.docx"                        # Save the document as a temporary file
    temp_file = "temp_doc.docx"
    doc.save(temp_file)
    doc_read = Document(temp_file)                     # Read the temporary file and check if the table exists
    tables = doc_read.tables
    assert len(tables) == 1                             # Check if the first table in the document matches the DataFrame shape
    table = tables[0]
    assert len(table.rows) == df.shape[0] + 1            # Include header row
    assert len(table.columns) == df.shape[1]+1           #include the result coloumn
    os.remove(temp_file) 


def test_graph(sample_dataframe):
    graph2 = polarity_graph(sample_dataframe)
    # Add assertions to check the generated graph

def test_pie(sample_dataframe):
    graph4 = polarity_pie(sample_dataframe)
    # Add assertions to check the generated graph