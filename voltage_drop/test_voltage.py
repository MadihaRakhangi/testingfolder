import pytest 
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import io
import os
from voltage import voltage_result, voltage_table, voltage_graph,voltage_pie

test_vf = pd.read_csv("voltage.csv")

@pytest.fixture
def sample_dataframe():
    data = {
        "Measured Voltage (V, L-N)[FROM]": [220, 210, 230, 240, 225, 215, 235, 225, 230, 220],
        "Measured Voltage (V, L-N)[TO]": [217, 208, 225, 235, 220, 212, 230, 218, 228, 218],
        "Type_ISS": ["Public", "Public", "Private", "Private", "Public", "Private", "Public", "Private", "Private", "Public"],
        "PoS": ["Lighting", "Other", "Lighting", "Other", "Other", "Lighting", "Other", "Lighting", "Other", "Lighting"],
        "Distance": [50, 80, 110, 90, 70, 120, 100, 60, 85, 95],
        # Add other required columns here
    }
    vf = pd.DataFrame(data)
    return vf


def test_voltage_table():
    vf = pd.DataFrame({                                                                                          #create  a test dataframe
        "Measured Voltage (V, L-N)[FROM]": [220, 210, 230, 240, 225, 215, 235, 225, 230, 220],
        "Measured Voltage (V, L-N)[TO]": [217, 208, 225, 235, 220, 212, 230, 218, 228, 218],
        "Type_ISS": ["Public", "Public", "Private", "Private", "Public", "Private", "Public", "Private", "Private", "Public"],
        "PoS": ["Lighting", "Other", "Lighting", "Other", "Other", "Lighting", "Other", "Lighting", "Other", "Lighting"],
        "Distance": [50, 80, 110, 90, 70, 120, 100, 60, 85, 95],
        # Add other required columns here
    })

    doc = Document()                                  # Create a Document object
    doc =voltage_table(vf, doc)                      
    temp_file = "temp_doc.docx"                        # Save the document as a temporary file
    temp_file = "temp_doc.docx"
    doc.save(temp_file)
    doc_read = Document(temp_file)                     # Read the temporary file and check if the table exists
    tables = doc_read.tables
    assert len(tables) == 1                             # Check if the first table in the document matches the DataFrame shape
    table = tables[0]
    assert len(table.rows) == vf.shape[0] + 1            # Include header row
    assert len(table.columns) == vf.shape[1]          #include the result coloumn
    os.remove(temp_file)
    
def test_voltage_result():
    # Test cases
    assert voltage_result(3, "Public", "Lighting", 0) == "Pass"
    assert voltage_result(4, "Public", "Other", 0) == "Fail"
    assert voltage_result(5, "Private", "Lighting", 0) == "Fail"
    assert voltage_result(7, "Private", "Other", 0) == "Pass"
    assert voltage_result(2, "Public", "Lighting", 50) == "Pass"
    assert voltage_result(3, "Public", "Other", 80) == "Fail"
    assert voltage_result(4, "Private", "Lighting", 110) == "Pass"
    assert voltage_result(6, "Private", "Other", 90) == "Pass"
    assert voltage_result(4, "Public", "Lighting", 120) == "Pass"
    assert voltage_result(6, "Public", "Other", 100) == "Pass"
    assert voltage_result(8, "Private", "Lighting", 60) == "Fail"
    assert voltage_result(9, "Private", "Other", 85) == "Fail"
    assert voltage_result(10, "Public", "Lighting", 95) == "Fail"