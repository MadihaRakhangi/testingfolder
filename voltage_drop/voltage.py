import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import io

V="voltage.csv"
vf = pd.read_csv("voltage.csv")



vf["Calculated Voltage Drop (V)"] = (
    vf["Measured Voltage (V, L-N)[FROM]"] - vf["Measured Voltage (V, L-N)[TO]"]
)
vf["Voltage Drop %"] = (
    vf["Calculated Voltage Drop (V)"] / vf["Measured Voltage (V, L-N)[FROM]"]
) * 100
vf["Voltage Drop %"] = vf["Voltage Drop %"].round(decimals=2)
vf.to_csv("voltage_upd.csv", index=False)

lim1 = 3
lim2 = 5
lim3 = 6
lim4 = 8



def voltage_result(VD, Type_ISS, PoS, Dist):
    if Dist <= 0:
        if VD <= lim1:
            if Type_ISS == "Public" and PoS == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif VD <= lim2:
            if Type_ISS == "Public" and PoS == "Other":
                return "Pass"
            else:
                return "Fail"
        elif VD <= lim3:
            if Type_ISS == "Private" and PoS == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif VD <= lim4:
            if Type_ISS == "Private" and PoS == "Other":
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"

    elif 0 < Dist <= 100:
        if VD <= (lim1 + (Dist * 0.005)):
            if Type_ISS == "Public" and PoS == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif VD <= (lim2 + (Dist * 0.005)):
            if Type_ISS == "Public" and PoS == "Other":
                return "Pass"
            else:
                return "Fail"
        elif VD <= (lim3 + (Dist * 0.005)):
            if Type_ISS == "Private" and PoS == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif VD <= (lim4 + (Dist * 0.005)):
            if Type_ISS == "Private" and PoS == "Other":
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"

    elif Dist > 100:
        if VD <= (lim1 + 0.5):
            if Type_ISS == "Public" and PoS == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif VD <= (lim2 + 0.5):
            if Type_ISS == "Public" and PoS == "Other":
                return "Pass"
            else:
                return "Fail"
        elif VD <= (lim3 + 0.5):
            if Type_ISS == "Private" and PoS == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif VD <= (lim4 + 0.5):
            if Type_ISS == "Private" and PoS == "Other":
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"



def voltage_rang(length):
    res = []
    for row in range(0, length):
        VD_val = vf.iloc[row, 12]
        Type_ISS_val = vf.iloc[row, 6]
        PoS_val = vf.iloc[row, 7]
        Dist = vf.iloc[row, 8] - 100  # Calculate Dist for each row
        res.append(voltage_result(VD_val, Type_ISS_val, PoS_val, Dist))
    return res


def voltage_table(vf, doc):
    table_data = vf.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.55,
        3: 0.54,
        4: 0.38,
        5: 0.56,
        6: 0.5,
        7: 0.48,
        8: 0.71,
        9: 0.43,
        10: 0.56,
        11: 0.56,
        12: 0.5,
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)
    Results = voltage_rang(num_rows)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(0.8)
    for i in range(num_rows):
        res_index = i
        table.cell(i + 1, num_cols).text = Results[res_index]
    font_size = 7

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc


def voltage_combined_graph(vf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = vf["Voltage Drop %"]
    y = vf["Calculated Voltage Drop (V)"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Voltage Drop %")
    plt.ylabel("Calculated Voltage Drop (V)")
    plt.title("Calculated Voltage Drop (V) VS Voltage Drop %")

    # Pie chart
    plt.subplot(122)
    vf['Result'] = voltage_rang(len(vf))  # Ensure you have the voltage_rang() function defined correctly
    vf_counts = vf['Result'].value_counts()
    labels = vf_counts.index.tolist()
    values = vf_counts.values.tolist()
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
    plt.axis('equal')
    plt.title('Test Results')

    # Save the combined graph
    plt.savefig("combined_voltage_graph.png")

    # Save the combined graph as bytes
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined




def main():
    doc = Document()
    doc.add_heading("RESISTANCE CONDUCTOR TEST", 0)
    doc = voltage_table(vf, doc)
    graph_combined = voltage_combined_graph(vf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))
    doc.save("voltage.docx")


main()