import pytest
import pandas as pd
from floortest import resistanceresult ,resistancerang, resistance_graph , resistance_pie , resistance_table
import io
import os
import matplotlib as plt
from docx import Document


test_df=pd.read_csv("floorfinal.csv")


@pytest.fixture
def sample_dataframe():
                                                           
    df = pd.DataFrame({                                           # Define a sample dataframe for testing
        "Nominal Voltage to Earth of System (V)": [230,230, 230],
        "Applied Test Voltage (V)": [230, 230,270],
        "EffectiveResistance": [50, 100, 50],
        "Distance from previous test location (m)":[1,1,1,]
    })
    return df


def test_result():
    assert resistanceresult(230, 230, 55,1) == "pass"
    assert resistanceresult(650, 650, 87,2) == "fail"

def test_rang():
    df = pd.DataFrame({                                  #create  a test dataframe
         "Nominal Voltage to Earth of System (V)": [230,230, 230],
         "Applied Test Voltage (V)": [230, 230,270],
         "EffectiveResistance": [50, 100, 50],
         "Distance from previous test location (m)":[1,1,1,]
    })
    assert resistancerang(df.shape[0]) == ["pass","fail", "pass"]

def test_create_table():
    df = pd.DataFrame({                               #create  a test dataframe
        "Serial number": [1, 2, 3],
        "Applied Test Voltage (V)": [230, 230,270],
        "EffectiveResistance": [50, 100, 50],
        "Distance from previous test location (m)":[1,1,1,]
    })

    doc = Document()                                  # Create a Document object
    doc = resistance_table(df, doc)                      
    temp_file = "temp_doc.docx"                        # Save the document as a temporary file
    temp_file = "temp_doc.docx"
    doc.save(temp_file)
    doc_read = Document(temp_file)                     # Read the temporary file and check if the table exists
    tables = doc_read.tables
    assert len(tables) == 1                             # Check if the first table in the document matches the DataFrame shape
    table = tables[0]
    assert len(table.rows) == df.shape[0] + 1            # Include header row
    assert len(table.columns) == df.shape[1]+1           #include the result coloumn
    os.remove(temp_file)    



def test_resistance_graph(sample_dataframe):
    graph1 = resistance_graph(sample_dataframe)
    # Add assertions to check the generated graph

def test_resistance_pie(sample_dataframe):
    graph3 = resistance_pie(sample_dataframe)
    # Add assertions to check the generated graph


