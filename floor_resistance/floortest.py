import pandas as pd
import matplotlib.pyplot as plt
import docx
from docx import Document
import csv
import io
import numpy as np
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

M = "floor.csv"
df = pd.read_csv(M)



df['Effective Resistance'] = df['Applied Test Voltage (V)'] / df['Measured Output Current (mA)']
df.to_csv('floorfinal.csv', index=False)



def resistanceresult(Nom_EVolt, ATV, Eff_Floor, Dist_loc):
    if Nom_EVolt <= 500 and Dist_loc >= 1:
        if ATV == Nom_EVolt and Eff_Floor >= 50:
            return "pass"
        else:
            return "fail"
    elif Nom_EVolt > 500 and Dist_loc >= 1:
        if ATV == Nom_EVolt and Eff_Floor >= 100:
            return "pass"
        else:
            return "fail"
    elif Dist_loc <= 1:
        return "fail"
    else:
        return "Invalid input"


def flooresistancerang(length):
    res = []
    for row in range(length):
        Nom_EVolt = df.iloc[row, 5]
        Dist_loc = df.iloc[row, 4]
        ATV = df.iloc[row, 6]
        Eff_Floor = df.iloc[row, 8]
        res.append(resistanceresult(Nom_EVolt, ATV, Eff_Floor, Dist_loc))
    return res


def flooresistance_table(df, doc):
    table_data = df.iloc[:, :]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.4,
        1: 0.6,
        2: 0.7,
        3: 0.65,
        4: 0.5,
        5: 0.5,
        6: 0.5,
        7: 0.7,
        8: 0.7,
        9: 0.65,
        10: 0.4,
        11: 0.8,
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        table.cell(0, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            if isinstance(value, float):
                value = "{:.2f}".format(value)
            table.cell(i, j).text = str(value)

    Results = flooresistancerang(num_rows)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(0.6)
    for i in range(num_rows):
        cell = table.cell(i + 1, num_cols)
        cell.text = Results[i]

        if Results[i] == "pass":
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="#5ac85a"/>'.format(nsdecls("w"))
            )  # Green color
            cell._tc.get_or_add_tcPr().append(shading_elm)
        elif Results[i] == "fail":
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="#dc0000"/>'.format(nsdecls("w"))
            )  # Red color
            cell._tc.get_or_add_tcPr().append(shading_elm)

    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    font_size = 8

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    return doc


def flooresistance_combined_graph(df):
    try:
        plt.figure(figsize=(16, 8))

        plt.subplot(121)  # Sort the DataFrame by "Location" in ascending order
        x = df["Location"]
        y = df["Effective Resistance"]
        colors = ["#b967ff", "#e0a899", "#fffb96", "#428bca"]  # Add more colors if needed
        sorted_indices = np.argsort(y)  # Sort the indices based on y values
        x_sorted = [x[i] for i in sorted_indices]
        y_sorted = [y[i] for i in sorted_indices]
        plt.bar(x_sorted, y_sorted, color=colors)
        plt.xlabel("Location")
        plt.ylabel("Effective Floor Resistance")
        plt.title("Location VS Effective Floor Resistance (Bar Graph)")


        # Pie chart
        plt.subplot(122)
        df['Result'] = flooresistancerang(df.shape[0])
        df_counts = df['Result'].value_counts()
        labels = df_counts.index.tolist()
        values = df_counts.values.tolist()
        colors = ["#5ac85a", "#dc0000"]
        plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
        plt.title('Test Results')
        plt.axis('equal')

        # Save the combined graph as bytes
        graph_combined1 = io.BytesIO()
        plt.savefig(graph_combined1)
        plt.close()

        # Check if the graph was saved successfully
        if graph_combined1.tell() == 0:
            print("Graph not found")
            return None

        return graph_combined1

    except Exception as e:
        # Handle the error
        print(f"An error occurred: {str(e)}")
        return None




def main():
    M = "floor.csv"
    df = pd.read_csv("floorfinal.csv")
    doc = docx.Document()
    doc.add_heading('FLOOR  AND WALL RESISTANCE TEST', 0)
    doc = flooresistance_table(df, doc)
    graph_combined = flooresistance_combined_graph(df)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(3))
    doc.save("floor.docx")


main()
