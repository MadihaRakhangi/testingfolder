import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import io
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

H= "func_ops.csv"
of = pd.read_csv(H)

def func_ops_result(func_chk,inter_chk):
    if func_chk == "OK" and inter_chk == "OK":
        return "pass"
    elif func_chk == "OK" and inter_chk == "Not OK":
        return "fail"
    elif func_chk == "Not OK " and inter_chk == "OK":
        return "fail"
    elif func_chk == "Not OK" and inter_chk == "OK":
        return "fail"
    else:
        return "Invalid"
    

def func_ops_rang(of):
    res6 = []
    for row in range(0, len(of)):
        func_chk = of.iloc[row]['Functional Check']
        inter_chk = of.iloc[row]['Interlock check']
        res6.append(func_ops_result(func_chk,inter_chk))
    return res6

def func_ops_table(of, doc):
    table_data = of.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False
    
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.55,
        3: 0.54,
        4: 0.38,
        5: 0.56,
        6: 0.5,
        7: 0.48,
        8: 0.71,
        9: 0.43,
        10: 0.56,
        11: 0.56,
        12: 0.5
    }
    
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)
    
    Results = func_ops_rang(of)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(0.8)
    
    for i in range(num_rows):
        cell = table.cell(i + 1, num_cols)
        cell.text = Results[i]
        
        if Results[i] == "pass":
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="#5ac85a"/>'.format(nsdecls("w"))
            )  # Green color
            cell._tc.get_or_add_tcPr().append(shading_elm)
        elif Results[i] == "fail":
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="#dc0000"/>'.format(nsdecls("w"))
            )  # Red color
            cell._tc.get_or_add_tcPr().append(shading_elm)
    
    font_size = 6.5
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    
    return doc


def func_ops_combined_graph(of):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = of["Interlock check"]
    y = of["SN"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Interlock check")
    plt.ylabel("SN")
    plt.title("SN VS Interlock check")

    # Pie chart
    plt.subplot(122)
    of['Result'] = func_ops_rang(of)
    of_counts = of['Result'].value_counts()
    labels = of_counts.index.tolist()
    values = of_counts.values.tolist()
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90,colors=colors)
    plt.axis('equal')
    plt.title('Test Results')
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined

def main():
    doc = Document()
    doc.add_heading('RESISTANCE CONDUCTOR TEST', 0)
    doc = func_ops_table(of, doc)
    graph_combined = func_ops_combined_graph(of)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))
    doc.save("funcandops.docx")


main()