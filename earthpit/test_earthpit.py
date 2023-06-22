import pytest
import pandas as pd
import io
import os
import matplotlib as plt
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

from earthpit import Earth_result, Earth_create_table, Earth_graph, Earth_Pie

test_ef=pd.read_csv("earthpit.csv")

@pytest.fixture
def sample_data():
    ef= {
        "Nearest Electrode Distance": [5, 2, 3, 4, 1],
        "Earth Electrode Depth": [5, 3, 2, 4, 1],
        "Measured Earth Resistance - Individual": [1, 2, 3, 4, 5],
        "No. of Parallel Electrodes": [1, 1, 1, 1, 1]
    }
    return pd.DataFrame(ef)

def test_Earth_result():
    assert Earth_result(1.0, 1) == "PASS - Test Electrodes are properly placed"
    assert Earth_result(0.5, 1) == "PASS - Test Electrodes are not properly placed"
    assert Earth_result(1.0, 3) == "FAIL - Test Electrodes are properly placed"
    assert Earth_result(0.5, 3) == "FAIL - Test Electrodes are not properly placed"
    assert Earth_result(1.0, 2) == "PASS - Test Electrodes are properly placed"

def test_create_table():
    ef = pd.DataFrame({                               #create  a test dataframe
        "Nearest Electrode Distance": [5, 2, 3, 4, 1],
        "Earth Electrode Depth": [5, 3, 2, 4, 1],
        "Measured Earth Resistance - Individual": [1, 2, 3, 4, 5],
        "No. of Parallel Electrodes": [1, 1, 1, 1, 1]
    })

    doc = Document()                                  # Create a Document object
    doc =Earth_create_table(ef, doc)                      
    temp_file = "temp_doc.docx"                        # Save the document as a temporary file
    temp_file = "temp_doc.docx"
    doc.save(temp_file)
    doc_read = Document(temp_file)                     # Read the temporary file and check if the table exists
    tables = doc_read.tables
    assert len(tables) == 1                             # Check if the first table in the document matches the DataFrame shape
    table = tables[0]
    assert len(table.rows) == ef.shape[0] + 1            # Include header row
    assert len(table.columns) == ef.shape[1]          #include the result coloumn
    os.remove(temp_file) 


def test_Earth_graph(sample_data):
    sample_data['Result'] = [
        "PASS - Test Electrodes are properly placed",
        "PASS - Test Electrodes are not properly placed",
        "FAIL - Test Electrodes are properly placed",
        "FAIL - Test Electrodes are not properly placed",
        "Invalid"
    ]

    graph = Earth_graph(sample_data)

    assert graph is not None
    # ... assert other properties of the graph


def test_Earth_Pie(sample_data):
    sample_data['Result'] = [
        "PASS - Test Electrodes are properly placed",
        "PASS - Test Electrodes are not properly placed",
        "FAIL - Test Electrodes are properly placed",
        "FAIL - Test Electrodes are not properly placed",
        "Invalid"
    ]
    
    graph = Earth_Pie(sample_data)

    assert graph is not None


# Run the tests
pytest.main()
