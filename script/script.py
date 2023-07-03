import pandas as pd
import matplotlib.pyplot as plt
import docx
from docx import Document
import csv
import io
from docx.shared import Inches
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
import numpy as np
import datetime
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

A = "floor.csv"                                                                   #all the csv files mentioned                                                 
df = pd.read_csv(A)
B = "Insulate.csv"
mf = pd.read_csv(B)
C = "phasesequence.csv"
pf = pd.read_csv("phasesequence.csv")
D = "pol.csv"
af = pd.read_csv("pol.csv")
E = "voltage.csv"
vf = pd.read_csv("voltage.csv")
F= "residual.csv"
rf = pd.read_csv("residual.csv")
G = "earthpit.csv"
ef = pd.read_csv(G)
H="threephase.csv"
tf=pd.read_csv("threephase.csv")
I="threephasevalue.csv"
tf2=pd.read_csv("threephasevalue.csv")
J="resistance.csv"
jf=pd.read_csv("resistance.csv")
H= "func_ops.csv"
of = pd.read_csv(H)
P="pat.csv"
bf = pd.read_csv("pat.csv")
Q="eli-socket.csv"
sf = pd.read_csv("eli-socket.csv")
R="sugg-max-eli.csv"
fs = pd.read_csv("sugg-max-eli.csv")
X="eli-test.csv"
gf = pd.read_csv("eli-test.csv")
Y="sugg-max-eli.csv"
fg = pd.read_csv("sugg-max-eli.csv")


df["EffectiveResistance"] = df["Applied Test Voltage (V)"] / df["Measured Output Current (mA)"]          #floor resistance calculation      
df.to_csv("floorfinal.csv", index=False)


vf["Calculated Voltage Drop (V)"] = (vf["Measured Voltage (V, L-N)[FROM]"] - vf["Measured Voltage (V, L-N)[TO]"])        #voltage drop calculation
vf["Voltage Drop %"] = (vf["Calculated Voltage Drop (V)"] / vf["Measured Voltage (V, L-N)[FROM]"]) * 100
vf["Voltage Drop %"] = vf["Voltage Drop %"].round(decimals=2)
vf.to_csv("voltage_upd.csv", index=False)

lim1 = 3
lim2 = 5
lim3 = 6
lim4 = 8


def alpha(Cond_T):                               # Function to calculate alpha value based on conductor typ      
    if Cond_T == "Al":
        return 0.0038
    elif Cond_T == "Cu":
        return 0.00429
    elif Cond_T == "GI":
        return 0.00641
    elif Cond_T == "SS":
        return 0.003
    else:
        return None

# Calculate Corrected Continuity Resistance
jf['Corrected Continuity Resistance (Ω)'] = jf['Continuity Resistance (?)'] - jf['Lead Internal Resistance (?)']

# Calculate Specific Conductor Resistance at 30°C
alpha_values = jf['Conductor Type'].apply(alpha)
jf['Specific Conductor Resistance (MΩ/m) at 30°C'] = jf['Corrected Continuity Resistance (Ω)'] / (
    1 + alpha_values * (jf['Conductor Temperature (°C)'] - 30))
jf['Specific Conductor Resistance (MΩ/m) at 30°C'] /= 1000000
jf['Specific Conductor Resistance (MΩ/m) at 30°C'] = jf['Specific Conductor Resistance (MΩ/m) at 30°C'].apply(
    lambda x: format(x, 'E'))
jf.to_csv('resistance_updated.csv', index=False)


sf1 = sf[
    [
        "SN",
        "Socket Name",
        "Location",
        "Facility Area",
        "Socket Type",
        "Earthing Configuration",
        "Upstream Breaker Rating (A)",
        "Upstream Breaker Make",
        "Upstream Breaker Type",
        "Trip Curve",
    ]
]

sf2 = sf[
    [
        "SN",
        "Socket Name",
        "Socket Rating (A)",
        "Socket Type",
        "No. of Phases",
        "V_LN",
        "V_LE",
        "V_NE",
        "L1-ELI",
        "L2-ELI",
        "L3-ELI",
        "Psc (kA)",
    ]
]

sf_filled = sf.fillna("")
sf["Upstream Breaker Rating (A)"] = sf["Upstream Breaker Rating (A)"].astype(int)

Device_Rating = sf["Upstream Breaker Rating (A)"]
No_phase = sf["No. of Phases"]
T_Curve = sf["Trip Curve"]
new_column = []
result_column = []
P = 0
K = 0
TMS = 1
TDS = 1

gf1 = gf[
    [
        "SN",
        "Device Name",
        "Location",
        "Facility Area",
        "Earthing Configuration",
        "Type of Circuit Location",
        "Device Rating (A)",
        "Device Make",
        "Device Type",
        "Device Sensitivity (mA)",
        "No. of Phases",
        "Trip Curve",
    ]
]

gf2 = gf[
    [
        "SN",
        "Device Name",
        "Device Rating (A)",
        "Device Type",
        "No. of Phases",
        "V_LN",
        "V_LE",
        "V_NE",
        "L1-ELI",
        "L2-ELI",
        "L3-ELI",
        "Psc (kA)",
    ]
]

gf_filled = gf.fillna("")
gf["Device Rating (A)"] = gf["Device Rating (A)"].astype(int)

Device_Rating = gf["Device Rating (A)"]
No_phase = gf["No. of Phases"]
T_Curve = gf["Trip Curve"]
new_column = []
result_column = []
P = 0
K = 0
TMS = 1
TDS = 1


def resistance_result(Nom_EVolt, ATV, Eff_Floor, Dist_loc):                                                     #floor and wall resistance result condition
    if Nom_EVolt <= 500 and Dist_loc >= 1:
        if ATV == Nom_EVolt and Eff_Floor >= 50:
            return "Pass"
        else:
            return "Fail"
    elif Nom_EVolt > 500 and Dist_loc >= 1:
        if ATV == Nom_EVolt and Eff_Floor >= 100:
            return "Pass"
        else:
            return "Fail"
    elif Dist_loc <= 1:
        return "Fail"
    else:
        return "Invalid input"


def insulate_result(Nom_CVolt, T_Volt, Insu_R):                                                                  #insulation result condition
    if Nom_CVolt <= 50:
        if Insu_R >= 0.5 and T_Volt == 250:
            return "Pass"
        else:
            return "Fail"
    elif 50 < Nom_CVolt <= 500:
        if Insu_R >= 1 and T_Volt == 500:
            return "Pass"
        else:
            return "Fail"
    elif Nom_CVolt > 500:
        if Insu_R >= 1 and T_Volt == 1000:
            return "Pass"
        else:
            return "Fail"
    else:
        return "Invalid input"


def phase_result(phase_seq):                                                                          #phase sequence result condition
    if phase_seq == "RYB":
        return "CLOCKWISE"
    else:
        return "ANTICLOCKWISE"


def polarity_result(line_neutral):                                                                      #polarity result condition
    if line_neutral == 230:
        return "OK"
    else:
        return "REVERSE"


def voltage_result(VD, Type_ISS, PoS, Dist):                                                             #voltage drop result condition
    if Dist <= 0:
        if VD <= lim1:
            if Type_ISS == "Public" and PoS == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif VD <= lim2:
            if Type_ISS == "Public" and PoS == "Other":
                return "Pass"
            else:
                return "Fail"
        elif VD <= lim3:
            if Type_ISS == "Private" and PoS == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif VD <= lim4:
            if Type_ISS == "Private" and PoS == "Other":
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"

    elif 0 < Dist <= 100:
        if VD <= (lim1 + (Dist * 0.005)):
            if Type_ISS == "Public" and PoS == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif VD <= (lim2 + (Dist * 0.005)):
            if Type_ISS == "Public" and PoS == "Other":
                return "Pass"
            else:
                return "Fail"
        elif VD <= (lim3 + (Dist * 0.005)):
            if Type_ISS == "Private" and PoS == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif VD <= (lim4 + (Dist * 0.005)):
            if Type_ISS == "Private" and PoS == "Other":
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"

    elif Dist > 100:
        if VD <= (lim1 + 0.5):
            if Type_ISS == "Public" and PoS == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif VD <= (lim2 + 0.5):
            if Type_ISS == "Public" and PoS == "Other":
                return "Pass"
            else:
                return "Fail"
        elif VD <= (lim3 + 0.5):
            if Type_ISS == "Private" and PoS == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif VD <= (lim4 + 0.5):
            if Type_ISS == "Private" and PoS == "Other":
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"


def resi_result(Type, Test_Current, Rated_OpCurrent, D_Tripped, Trip_Time):                                    #residual test result condition
    if Type == "AC":
        if Test_Current == 0.5 * Rated_OpCurrent:
                if D_Tripped == "No":
                    return "Pass"
                else:
                    return "Fail"
        elif Test_Current == 1 * Rated_OpCurrent and D_Tripped == "Yes":
            if Trip_Time <= 300:
                return "Pass"
            else:
                return "Fail"
        elif Test_Current == 2 * Rated_OpCurrent and D_Tripped == "Yes":
            if Trip_Time <= 150:
                return "Pass"
            else:
                return "Fail"
        elif Test_Current == 5 * Rated_OpCurrent and D_Tripped == "Yes":
            if Trip_Time <= 40:
                return "Pass"
            else:
                return "Fail"
        else:
            return "/="
    elif Type == "A":
        if Test_Current == 0.5 * Rated_OpCurrent:
            if D_Tripped == "No":
                return "Pass"
            else:
                return "Fail"
        elif Test_Current == 1 * Rated_OpCurrent and D_Tripped == "Yes":
            if 130 <= Trip_Time <= 500:
                return "Pass"
            else:
                return "Fail"
        elif Test_Current == 2 * Rated_OpCurrent and D_Tripped == "Yes":
            if 60 <= Trip_Time <= 200:
                return "Pass"
            else:
                return "Fail"
        elif Test_Current == 5 * Rated_OpCurrent and D_Tripped == "Yes":
            if 50 <= Trip_Time <= 150:
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"
    else:
        return "Pass"


def Earth_result(Elec_DistRatio, Mea_EarthResist):
    if Mea_EarthResist <= 2 and Elec_DistRatio >= 1:
        return "PASS - Test Electrodes are properly placed"
    elif Mea_EarthResist <= 2 and Elec_DistRatio < 1:
        return "PASS - Test Electrodes are not properly placed"
    elif Mea_EarthResist > 2 and Elec_DistRatio >= 1:
        return "FAIL - Test Electrodes are properly placed"
    elif Mea_EarthResist > 2 and Elec_DistRatio < 1:
        return "FAIL - Test Electrodes are not properly placed"
    else:
        return "Invalid"
    
def earthpit_result(Elec_DistRatio, Mea_EarthResist):                                                      #earth residual test condition
    if Mea_EarthResist <= 2 and Elec_DistRatio >= 1:
        return "PASS"
    elif Mea_EarthResist <= 2 and Elec_DistRatio < 1:
        return "PASS"
    elif Mea_EarthResist > 2 and Elec_DistRatio >= 1:
        return "FAIL"
    elif Mea_EarthResist > 2 and Elec_DistRatio < 1:
        return "FAIL"
    else:
        return "Invalid"
    
def earth_remark_result(Elec_DistRatio, Mea_EarthResist):                                                      #earth residual test condition
    if Mea_EarthResist <= 2 and Elec_DistRatio >= 1:
        return "Test Electrodes are properly placed"
    elif Mea_EarthResist <= 2 and Elec_DistRatio < 1:
        return "Test Electrodes are not properly placed"
    elif Mea_EarthResist > 2 and Elec_DistRatio >= 1:
        return "Test Electrodes are properly placed"
    elif Mea_EarthResist > 2 and Elec_DistRatio < 1:
        return "Test Electrodes are not properly placed"
    else:
        return "Invalid"
    
def threephase_result(tf, tf2):                                                                             #three phase symmetry result condition
    tf["Rated Line Voltage (V)"] = tf2["Rated Line Voltage (V)"]
    tf["Average Line Voltage (V)"] = round(
        (tf2["Voltage-L1L2 (V)"] + tf2["Voltage-L2L3 (V)"] + tf2["Voltage-L3L1 (V)"]) / 3, 2
    )

    tf["Average Phase Voltage (V)"] = (
        tf2["Voltage-L1N (V)"] + tf2["Voltage-L2N (V)"] + tf2["Voltage-L3N (V)"]
    ) / 3
    tf["Voltage Unbalance %"] = round(
        (
            (tf2["Voltage-L1N (V)"] - tf["Average Line Voltage (V)"])
            .abs()
            .where(
                (tf2["Voltage-L1N (V)"] - tf["Average Line Voltage (V)"]) > 0,
                (tf["Average Line Voltage (V)"] - tf2["Voltage-L1N (V)"]).abs(),
            )
            .max(axis=0)
            / tf["Average Line Voltage (V)"]
        )
        * 100,
        2,
    )
    tf["Voltage Result"] = np.where(tf["Voltage Unbalance %"] <= 10, "PASS", "FAIL")
    tf["Rated Phase Current (A)"] = tf2["Rated Phase Current (A)"]
    tf["Average Phase Current (A)"] = round(
        (tf2["Current-L1 (A)"] + tf2["Current-L2 (A)"] + tf2["Current-L3 (A)"]) / 3, 2
    )
    tf["Current Unbalance %"] = round(
        (
            (tf2["Current-L1 (A)"] - tf["Average Phase Current (A)"])
            .abs()
            .where(
                (tf2["Current-L1 (A)"] - tf["Average Phase Current (A)"]) > 0,
                (tf["Average Phase Current (A)"] - tf2["Current-L1 (A)"]).abs(),
            )
            .max(axis=0)
            / tf["Average Line Voltage (V)"]
        )
        * 100,
        2,
    )
    tf["Current Result"] = np.where(tf["Current Unbalance %"] <= 10, "PASS", "FAIL")
    tf["Voltage-NE (V)"] = tf2["Voltage-NE (V)"]
    tf["NEV Result"] = np.where(tf["Voltage-NE (V)"] <= 2, "PASS", "FAIL")
    tf["Zero Sum Current (mA)"] = tf2["Zero Sum Current (mA)"]
    tf["ZeroSum Result"] = np.where(
        tf["Zero Sum Current (mA)"] <= (tf["Rated Phase Current (A)"] / 5000 * 1000),
        "PASS",
        "FAIL",
    )
    return tf

def resc_result(Conti, C_ContR):
    if C_ContR <= 1:
        if Conti == "Yes":
            return "Pass"
        elif Conti == "No":
            return "Check Again"
        else:
            return "Invalid"
    elif C_ContR > 1:
        if Conti == "Yes":
            return "Fail"
        else:
            return "Fail"
        

def func_ops_result(func_chk,inter_chk):
    if func_chk == "OK" and inter_chk == "OK":
        return "pass"
    elif func_chk == "OK" and inter_chk == "Not OK":
        return "fail"
    elif func_chk == "Not OK " and inter_chk == "OK":
        return "fail"
    elif func_chk == "Not OK" and inter_chk == "OK":
        return "fail"
    else:
        return "Invalid"
    

def socket_result1(sf1, df2):
    I = Is * ((((K * TMS) / Td) + 1) ** (1 / P))
    if row["Earthing Configuration"] == "TN":
        IEC_val_TN = row["V_LE"] / I
        new_column.append(round(IEC_val_TN, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= IEC_val_TN
                and row["L2-ELI"] <= IEC_val_TN
                and row["L3-ELI"] <= IEC_val_TN
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= IEC_val_TN:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")
    elif row["Earthing Configuration"] == "TT":
        IEC_val_TT = 50 / I
        new_column.append(round(IEC_val_TT, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= IEC_val_TT
                and row["L2-ELI"] <= IEC_val_TT
                and row["L3-ELI"] <= IEC_val_TT
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= IEC_val_TT:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")


def socket_result2(sf1, df2):
    I = Is * (((((A / ((Td / TDS) - B)) + 1)) ** (1 / p)))
    if row["Earthing Configuration"] == "TN":
        IEEE_val_TN = row["V_LE"] / I
        new_column.append(round(IEEE_val_TN, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= IEEE_val_TN
                and row["L2-ELI"] <= IEEE_val_TN
                and row["L3-ELI"] <= IEEE_val_TN
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= IEEE_val_TN:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")
    elif row["Earthing Configuration"] == "TT":
        IEEE_val_TT = 50 / I
        new_column.append(round(IEEE_val_TT, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= IEEE_val_TT
                and row["L2-ELI"] <= IEEE_val_TT
                and row["L3-ELI"] <= IEEE_val_TT
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= IEEE_val_TT:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")


for index, row in sf.iterrows():
    if row["Upstream Breaker Type"] == "MCB":
        rating = row[8]
        trip = row[13]  # Assuming the "Trip Curve" column is at index 10
        result_row = fs[fs["Device Rating (A)"] == rating]
        if trip in result_row.columns:
            val_MCB = result_row[trip].values[0]
        else:
            val_MCB = (
                0  # Set a default value when 'Trip Curve' value is not found in sugg-max-eli.csv
            )
        new_column.append(round(val_MCB, 2))

        if row["No. of Phases"] == 3:
            if row["L1-ELI"] <= val_MCB and row["L2-ELI"] <= val_MCB and row["L3-ELI"] <= val_MCB:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= val_MCB:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")

    elif (
        row["Upstream Breaker Type"] in ["RCD", "RCBO", "RCCB"]
        and row["Earthing Configuration"] == "TN"
    ):
        rccb_val_TN = (row["V_LE"] / row["Upstream Breaker Sensitivity"]) * 1000
        new_column.append(round(rccb_val_TN, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= rccb_val_TN
                and row["L2-ELI"] <= rccb_val_TN
                and row["L3-ELI"] <= rccb_val_TN
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= rccb_val_TN:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")

    elif (
        row["Upstream Breaker Type"] in ["RCD", "RCBO", "RCCB"]
        and row["Earthing Configuration"] == "TT"
    ):
        rccb_val_TT = (50 / row["Upstream Breaker Sensitivity"]) * 1000
        new_column.append(round(rccb_val_TT, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= rccb_val_TT
                and row["L2-ELI"] <= rccb_val_TT
                and row["L3-ELI"] <= rccb_val_TT
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= rccb_val_TT:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")

    elif row["Upstream Breaker Type"] == "MCCB":
        if row["Type of Circuit Location"] == "Final":
            Td = 0.4
        elif row["Type of Circuit Location"] == "Distribution":
            Td = 5
        Is = row["Upstream Breaker Rating (A)"]
        if row[13] == "IEC Standard Inverse":
            P = 0.02
            K = 0.14
            socket_result1(sf1, sf2)
        elif row[13] == "IEC Very Inverse":
            P = 1
            K = 13.5
            socket_result1(sf1, sf2)
        elif row[13] == "IEC Long-Time Inverse":
            P = 1
            K = 120
            socket_result1(sf1, sf2)
        elif row[13] == "IEC Extremely Inverse":
            P = 2
            K = 80
            socket_result1(sf1, sf2)
        elif row[13] == "IEC Ultra Inverse":
            P = 2.5
            K = 315.2
            socket_result1(sf1, sf2)
        elif row[13] == "IEEE Moderately Inverse":
            A = 0.0515
            B = 0.114
            p = 0.02
            socket_result1(sf1, sf2)
        elif row[13] == "IEEE Very Inverse":
            A = 19.61
            B = 0.491
            p = 2
            socket_result1(sf1, sf2)
        elif row[13] == "IEEE Extremely Inverse":
            A = 28.2
            B = 0.1217
            p = 2
            socket_result2(sf1, sf2)


new_column = pd.Series(new_column[: len(sf2)], name="Suggested Max ELI (Ω)")
sf2.loc[:, "Suggested Max ELI (Ω)"] = new_column
sf2.loc[:, "Suggested Max ELI (Ω)"] = sf2["Suggested Max ELI (Ω)"].apply(lambda x: "{:.2f}".format(x))
result_column = pd.Series(result_column[: len(sf2)], name="Result")
sf2["Result"] = result_column

def eli_test_result1(gf1, gf2):
    I = Is * ((((K * TMS) / Td) + 1) ** (1 / P))
    if row["Earthing Configuration"] == "TN":
        IEC_val_TN = row["V_LE"] / I
        new_column.append(round(IEC_val_TN, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= IEC_val_TN
                and row["L2-ELI"] <= IEC_val_TN
                and row["L3-ELI"] <= IEC_val_TN
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= IEC_val_TN:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")
    elif row["Earthing Configuration"] == "TT":
        IEC_val_TT = 50 / I
        new_column.append(round(IEC_val_TT, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= IEC_val_TT
                and row["L2-ELI"] <= IEC_val_TT
                and row["L3-ELI"] <= IEC_val_TT
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= IEC_val_TT:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")


def eli_test_result2(gf1, gf2):
    I = Is * (((((A / ((Td / TDS) - B)) + 1)) ** (1 / p)))
    if row["Earthing Configuration"] == "TN":
        IEEE_val_TN = row["V_LE"] / I
        new_column.append(round(IEEE_val_TN, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= IEEE_val_TN
                and row["L2-ELI"] <= IEEE_val_TN
                and row["L3-ELI"] <= IEEE_val_TN
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= IEEE_val_TN:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")
    elif row["Earthing Configuration"] == "TT":
        IEEE_val_TT = 50 / I
        new_column.append(round(IEEE_val_TT, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= IEEE_val_TT
                and row["L2-ELI"] <= IEEE_val_TT
                and row["L3-ELI"] <= IEEE_val_TT
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= IEEE_val_TT:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")


for index, row in gf.iterrows():
    if row["Device Type"] == "MCB":
        rating = row[6]
        trip = row[11]  # Assuming the "Trip Curve" column is at index 10
        result_row = fg[fg["Device Rating (A)"] == rating]
        if trip in result_row.columns:
            val_MCB = result_row[trip].values[0]
        else:
            val_MCB = (
                0  # Set a default value when 'Trip Curve' value is not found in sugg-max-eli.csv
            )
        new_column.append(round(val_MCB, 2))

        if row["No. of Phases"] == 3:
            if row["L1-ELI"] <= val_MCB and row["L2-ELI"] <= val_MCB and row["L3-ELI"] <= val_MCB:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= val_MCB:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")

    elif row["Device Type"] in ["RCD", "RCBO", "RCCB"] and row["Earthing Configuration"] == "TN":
        rccb_val_TN = (row["V_LE"] / row["Device Sensitivity (mA)"]) * 1000
        new_column.append(round(rccb_val_TN, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= rccb_val_TN
                and row["L2-ELI"] <= rccb_val_TN
                and row["L3-ELI"] <= rccb_val_TN
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= rccb_val_TN:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")

    elif row["Device Type"] in ["RCD", "RCBO", "RCCB"] and row["Earthing Configuration"] == "TT":
        rccb_val_TT = (50 / row["Device Sensitivity (mA)"]) * 1000
        new_column.append(round(rccb_val_TT, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= rccb_val_TT
                and row["L2-ELI"] <= rccb_val_TT
                and row["L3-ELI"] <= rccb_val_TT
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= rccb_val_TT:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")

    elif row["Device Type"] == "MCCB" or row["Device Type"] == "ACB":
        if row["Type of Circuit Location"] == "Final":
            Td = 0.4
        elif row["Type of Circuit Location"] == "Distribution":
            Td = 5
        Is = row["Device Rating (A)"]
        if row[11] == "IEC Standard Inverse":
            P = 0.02
            K = 0.14
            eli_test_result1(gf1, gf2)
        elif row[11] == "IEC Very Inverse":
            P = 1
            K = 13.5
            eli_test_result1(gf1, gf2)
        elif row[11] == "IEC Long-Time Inverse":
            P = 1
            K = 120
            eli_test_result1(gf1, gf2)
        elif row[11] == "IEC Extremely Inverse":
            P = 2
            K = 80
            eli_test_result1(gf1, gf2)
        elif row[11] == "IEC Ultra Inverse":
            P = 2.5
            K = 315.2
            eli_test_result1(gf1, gf2)
        elif row[11] == "IEEE Moderately Inverse":
            A = 0.0515
            B = 0.114
            p = 0.02
            eli_test_result2(gf1, gf2)
        elif row[11] == "IEEE Very Inverse":
            A = 19.61
            B = 0.491
            p = 2
            eli_test_result2(gf1, gf2)
        elif row[11] == "IEEE Extremely Inverse":
            A = 28.2
            B = 0.1217
            p = 2
            eli_test_result2(gf1, gf2)

new_column = pd.Series(new_column[: len(gf2)], name="Suggested Max ELI (Ω)")
gf2["Suggested Max ELI (Ω)"] = new_column
gf2["Suggested Max ELI (Ω)"] = gf2["Suggested Max ELI (Ω)"].apply(lambda x: "{:.2f}".format(x))
result_column = pd.Series(result_column[: len(gf2)], name="Result")
gf2["Result"] = result_column


        
def flooresistance_rang(length):                                                                     #gives  floor and wall resistance result coloumn
    res1 = []
    for row in range(length):
        Nom_EVolt = df.iloc[row, 5]
        Dist_loc = df.iloc[row, 4]
        ATV = df.iloc[row, 6]
        Eff_Floor = df.iloc[row, 8]
        res1.append(resistance_result(Nom_EVolt, ATV, Eff_Floor, Dist_loc))
    return res1


def insulation_rang(length):                                                                   #gives insulation result coloumn
    res2 = []
    for row in range(length):  # Adjusted the range to start from 0
        Nom_CVolt = mf.iloc[row, 7]
        T_Volt = mf.iloc[row, 9]
        Insu_R = mf.iloc[row, 13]
        res2.append(insulate_result(Nom_CVolt, T_Volt, Insu_R))
        print(Nom_CVolt)
    return res2


def phase_rang(pf):                                                                              #gives phase sequence result coloumn
    res3 = []
    phase_seqs = pf["Phase Sequence"]
    for seq in phase_seqs:
        if seq == "RYB":
            res3.append("CLOCKWISE")
        elif seq == "RBY":
            res3.append("ANTICLOCKWISE")
        else:
            res3.append("UNKNOWN")
    return res3


def polarity_rang(length):                                                                         #gives polarity result coloumn
    res4 = []
    for row in range(0, length):
        line_neutral = af.iloc[row, 5]
        res4.append(polarity_result(line_neutral))
    return res4


def voltage_rang(length):                                                                         #gives voltage result coloumn
    res = []
    for row in range(0, length):
        VD_val = vf.iloc[row, 12]
        Type_ISS_val = vf.iloc[row, 6]
        PoS_val = vf.iloc[row, 7]
        Dist = vf.iloc[row, 8] - 100  # Calculate Dist for each row
        res.append(voltage_result(VD_val, Type_ISS_val, PoS_val, Dist))
    return res

def earthpit_rang(length):
    res6 = []
    for row in range(length):  # Adjusted the range to start from 0
        Elec_DistRatio= ef.iloc[row, 7]
        Mea_EarthResist = ef.iloc[row, 8]
        res6.append(earthpit_result(Elec_DistRatio, Mea_EarthResist))
    return res6



def resc_rang(jf):
    res5 = []
    for row in range(0, len(jf)):
        Conti = jf.iloc[row]['Is Continuity found?']
        C_ContR = jf.iloc[row]['Corrected Continuity Resistance (Ω)']
        res5.append(resc_result(Conti, C_ContR))
    return res5

def func_ops_rang(of):
    res6 = []
    for row in range(0, len(of)):
        func_chk = of.iloc[row]['Functional Check']
        inter_chk = of.iloc[row]['Interlock check']
        res6.append(func_ops_result(func_chk,inter_chk))
    return res6


def resistance_table(df, doc):                                                                      #creates the floor and wall resistance table with  result coloumn
    table_data = df.iloc[:, :]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.4,
        1: 0.6,
        2: 0.7,
        3: 0.65,
        4: 0.5,
        5: 0.5,
        6: 0.5,
        7: 0.5,
        8: 0.5,
        9: 0.5,
        10: 0.7,
        11: 0.6,
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    first_row_cells = table.rows[0].cells
    for cell in first_row_cells:
        cell_elem = cell._element
        tc_pr = cell_elem.get_or_add_tcPr()
        shading_elem = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>'
        )
        tc_pr.append(shading_elem)
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            if isinstance(value, float):
                value = "{:.2f}".format(value)
            table.cell(i, j).text = str(value)

    Results = flooresistance_rang(num_rows)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(0.6)
    for i in range(num_rows):
        table.cell(i + 1, num_cols).text = Results[i]
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    font_size = 8

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    return doc


def insulation_table(df, doc):                                                                #creates the insulation table with  result coloumn
    table_data = df.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)  # Add +1 for the "Result" column
    table.style = "Table Grid"
    table.autofit = False
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.55,
        3: 0.54,
        4: 0.38,
        5: 0.4,
        6: 0.4,
        7: 0.48,
        8: 0.5,
        9: 0.3,
        10: 0.5,
        11: 0.4,
        12: 0.4,
        13: 0.4,
        num_cols: 0.8,  # Width for the "Result" column
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    first_row_cells = table.rows[0].cells
    for cell in first_row_cells:
        cell_elem = cell._element
        tc_pr = cell_elem.get_or_add_tcPr()
        shading_elem = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>'
        )
        tc_pr.append(shading_elem)
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)
    Results = insulation_rang(num_rows)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(
        column_widths[num_cols]
    )  # Set width for the "Result" column
    for i in range(0, num_rows):
        res_index = i - 1
        table.cell(i + 1, num_cols).text = Results[res_index]
        if Results == "Pass":
            shading_elem = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>'
                )
    font_size = 6.5

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    return doc


def phase_table(df, doc):                                                                                #creates the phase sequence table with  result coloumn
    table_data = df.iloc[:, :]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.2,
        1: 0.4,
        2: 0.4,
        3: 0.4,
        4: 0.6,
        5: 0.3,
        6: 0.3,
        7: 0.3,
        8: 0.3,
        9: 0.3,
        10: 0.3,
        11: 0.3,
        12: 0.6,
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    first_row_cells = table.rows[0].cells
    for cell in first_row_cells:
        cell_elem = cell._element
        tc_pr = cell_elem.get_or_add_tcPr()
        shading_elem = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>'
        )
        tc_pr.append(shading_elem)
    for i, row in enumerate(table_data.itertuples(), start=0):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i + 1, j).text = str(value)
    results = phase_rang(pf)

    table.cell(0, num_cols).text = "Result"
    for i, result in enumerate(results, start=1):
        table.cell(i, num_cols).text = result
        table.cell(i, num_cols).width = Inches(0.8)

    font_size = 8
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc


def polarity_table(af, doc):                                              #creates the polairty table with  result coloumn
    table_data = af.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.55,
        3: 0.54,
        4: 0.38,
        5: 0.56,
        6: 0.5,
        7: 0.48,
        8: 0.71,
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    first_row_cells = table.rows[0].cells
    for cell in first_row_cells:
        cell_elem = cell._element
        tc_pr = cell_elem.get_or_add_tcPr()
        shading_elem = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>'
        )
        tc_pr.append(shading_elem)
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)
    Results = polarity_rang(num_rows)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(0.8)
    for i in range(num_rows):
        res_index = i
        table.cell(i + 1, num_cols).text = Results[res_index]
    font_size = 8

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc


def voltage_table(vf, doc):                                                                      #creates the voltage table with  result coloumn
    table_data = vf.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.55,
        3: 0.4,
        4: 0.4,
        5: 0.4,
        6: 0.5,
        7: 0.48,
        8: 0.5,
        9: 0.43,
        10: 0.56,
        11: 0.56,
        12: 0.5,
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    first_row_cells = table.rows[0].cells
    for cell in first_row_cells:
        cell_elem = cell._element
        tc_pr = cell_elem.get_or_add_tcPr()
        shading_elem = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>'
        )
        tc_pr.append(shading_elem)
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)
    Results = voltage_rang(num_rows)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(0.8)
    for i in range(num_rows):
        res_index = i
        table.cell(i + 1, num_cols).text = Results[res_index]
    font_size = 7

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc

def residual_table(rf, doc):                                                             #creates the residual  table with  result coloumn
    rf["Result"] = np.NaN
    for index, row in rf.iterrows():
        trip_time_str = row["Trip Time (ms)"]
        if trip_time_str.isnumeric():
            trip_time = int(trip_time_str)
            result_val = resi_result(
                row["Type"],
                row["Test Current (mA)"],
                row["Rated Residual Operating Current,I?n (mA)"],
                row["Device Tripped"],
                trip_time,
            )
            rf.loc[index, "Result"] = result_val
        elif trip_time_str == "-":
            rf.loc[index, "Result"] = "Pass"
        else:
            rf.loc[index, "Result"] = "invalid"

    table_data = rf.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    table.autofit = False
    column_widths = {
        0: 0.3,
        1: 0.5,
        2: 0.5,
        3: 0.3,
        4: 0.3,
        5: 0.3,
        6: 0.3,
        7: 0.3,
        8: 0.37,
        9: 0.3,
        10: 0.3,
        11: 0.3,
        12: 0.48,
        13: 0.4,
        14: 0.49,
        15: 0.5,
        16: 0.4,  # Width for Result column
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    first_row_cells = table.rows[0].cells
    for cell in first_row_cells:
        cell_elem = cell._element
        tc_pr = cell_elem.get_or_add_tcPr()
        shading_elem = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>'
        )
        tc_pr.append(shading_elem)
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)

    font_size = 5
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc


def earthpit_table(ef, doc):                                                                      #creates the earthpit table with  result coloumn
    ef["Electrode Distance Ratio"] = round(
        ef["Nearest Electrode Distance"] / ef["Earth Electrode Depth"], 2
    )
    ef["Calculated Earth Resistance - Individual (Ω)"] = (
        ef["Measured Earth Resistance - Individual"] * ef["No. of Parallel Electrodes"]
    )

    ef["Remark"] = ef.apply(
        lambda row: earth_remark_result(
            row["Electrode Distance Ratio"], row["Measured Earth Resistance - Individual"]
        ),
        axis=1,
    )
    ef["Result"] = ef.apply(
        lambda row: earthpit_result(
            row["Electrode Distance Ratio"], row["Measured Earth Resistance - Individual"]
        ),
        axis=1,
    )

    table_data = ef.iloc[:, 0:]
    num_rows, num_cols = table_data.shape[0], table_data.shape[1]
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    table.autofit = False

    column_widths = {
        0: 0.25,
        1: 0.54,
        2: 0.6,
        3: 0.5,
        4: 0.48,
        5: 0.56,
        6: 0.4,
        7: 0.4,
        8: 0.4,
        9: 0.55,
        10: 0.5,
        11: 0.6,
        12: 0.5
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths.get(j, 1))
    first_row_cells = table.rows[0].cells
    for cell in first_row_cells:
        cell_elem = cell._element
        tc_pr = cell_elem.get_or_add_tcPr()
        shading_elem = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>'
        )
        tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    return doc


def threephase_table(tf, doc):                                                                 #creates the three phase table with  result coloumn
    table_data = tf.values
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    table.autofit = False
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.5,
        3: 0.54,
        4: 0.42,
        5: 0.4,
        6: 0.4,
        7: 0.4,
        8: 0.46,
        9: 0.43,
        10: 0.4,
        11: 0.4,
        12: 0.4,
        13: 0.4,
        14: 0.4,
        15: 0.4,
        16: 0.4,
    }
    for j, col in enumerate(tf.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    first_row_cells = table.rows[0].cells
    for cell in first_row_cells:
        cell_elem = cell._element
        tc_pr = cell_elem.get_or_add_tcPr()
        shading_elem = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>'
        )
        tc_pr.append(shading_elem)   
    for i, row in enumerate(table_data, start=1):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)
        font_size = 6
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc

def resc_table(jf, doc):
    table_data = jf.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.4,
        3: 0.4,
        4: 0.4,
        5: 0.4,
        6: 0.4,
        7: 0.4,
        8: 0.71,
        9: 0.43,
        10: 0.4,
        11: 0.56, 
        12:0.5   # Add this line
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    first_row_cells = table.rows[0].cells
    for cell in first_row_cells:
        cell_elem = cell._element
        tc_pr = cell_elem.get_or_add_tcPr()
        shading_elem = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>'
        )
        tc_pr.append(shading_elem) 
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)
    Results = resc_rang(jf)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(0.8)
    for i in range(num_rows):
        res_index = i
        table.cell(i + 1, num_cols).text = Results[res_index]
    font_size = 6.5

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc

def func_ops_table(of, doc):
    table_data = of.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.55,
        3: 0.54,
        4: 0.38,
        5: 0.56,
        6: 0.5,
        7: 0.48,
        8: 0.71,
        9: 0.43,
        10: 0.56,
        11: 0.56, 
        12:0.5   # Add this line
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    first_row_cells = table.rows[0].cells
    for cell in first_row_cells:
        cell_elem = cell._element
        tc_pr = cell_elem.get_or_add_tcPr()
        shading_elem = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>'
        )
        tc_pr.append(shading_elem) 
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)
    Results = func_ops_rang(of)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(0.8)
    for i in range(num_rows):
        res_index = i
        table.cell(i + 1, num_cols).text = Results[res_index]
    font_size = 7

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc

def pat_table(bf, doc):
    table_data = bf.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows+1 , cols=num_cols )
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.2,
        1: 0.5,
        2: 0.6,
        3: 0.6,
        4: 0.8,
        5: 0.7,
        6: 0.6,
        7: 0.6,
        8: 0.6,
        9: 0.7,
        10: 0.48,
        11: 0.6,
        12: 0.9,
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths.get(j, 1))
        table.cell(0, j).paragraphs[
            0
        ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.paragraphs[
                0
            ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center
            if j == num_cols - 1:  # Apply background color only to the Result column
                result_cell = cell
                if value == "PASS":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="#5ac85a"/>'.format(nsdecls("w"))
                    )  # Green color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
                elif value == "FAIL":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="#dc0000"/>'.format(nsdecls("w"))
                    )  # Red color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"

    return doc


def create_eli_table1(sf1, doc):
    sf1 = sf1.fillna("")
    table_data = df.iloc[:, :]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols )
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.2,
        1: 0.5,
        2: 0.75,
        3: 0.6,
        4: 0.8,
        5: 0.8,
        6: 0.6,
        7: 0.8,
        8: 1,
        9: 1.25,
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths.get(j, 1))
        table.cell(0, j).paragraphs[
            0
        ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.paragraphs[
                0
            ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center
            if j == num_cols - 1:  # Apply background color only to the Result column
                result_cell = cell
                if value == "Pass":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls("w"))
                    )  # Green color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
                elif value == "Fail":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls("w"))
                    )  # Red color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"

    return doc


def create_eli_table2(sf2, doc):
    table_data = sf2.iloc[:, :]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols )
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.2,
        1: 0.44,
        2: 0.55,
        3: 0.59,
        4: 0.6,
        5: 0.58,
        6: 0.45,
        7: 0.47,
        8: 0.45,
        9: 0.41,
        10: 0.41,
        11: 0.41,
        12: 0.62,
        13: 0.8,
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths.get(j, 1))
        table.cell(0, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        table.cell(0, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if j == num_cols - 1:  # Apply background color only to the Result column
                result_cell = cell
                if value == "Pass":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls("w"))
                    )  # Green color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
                elif value == "Fail":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls("w"))
                    )  # Red color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"

    return doc

def eli_test_table1(gf1, doc):
    gf1 = gf1.fillna("")
    doc.add_heading("Earth Loop Impedance Test - Circuit Breaker", level=1)
    table_data = gf1.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols )
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.2,
        1: 0.5,
        2: 0.6,
        3: 0.6,
        4: 0.8,
        5: 0.7,
        6: 0.6,
        7: 0.6,
        8: 0.55,
        9: 0.7,
        10: 0.41,
        11: 1.2,
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths.get(j, 1))
        table.cell(0, j).paragraphs[
            0
        ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.paragraphs[
                0
            ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center
            if j == num_cols - 1:  # Apply background color only to the Result column
                result_cell = cell
                if value == "Pass":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls("w"))
                    )  # Green color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
                elif value == "Fail":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls("w"))
                    )  # Red color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"

    return doc


def eli_test_table2(df2, doc):
    table_data = df2.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols )
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.2,
        1: 0.44,
        2: 0.55,
        3: 0.59,
        4: 0.6,
        5: 0.58,
        6: 0.45,
        7: 0.47,
        8: 0.45,
        9: 0.41,
        10: 0.41,
        11: 0.41,
        12: 0.6,
        13: 0.9,
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths.get(j, 1))
        table.cell(0, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if j == num_cols - 1:  # Apply background color only to the Result column
                result_cell = cell
                if value == "Pass":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls("w"))
                    )  # Green color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
                elif value == "Fail":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls("w"))
                    )  # Red color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"

    return doc


def flooresistance_combined_graph(df):
    plt.figure(figsize=(16, 8))

    # bar graph
    plt.subplot(121)
    x = df["Location"]
    y = df["EffectiveResistance"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Location")
    plt.ylabel("Effectivefloor")
    plt.title("Location VS Effectivefloor (Scatter Plot)")

    # Pie chart
    plt.subplot(122)
    df['Result'] =flooresistance_rang(df.shape[0])
    df_counts = df['Result'].value_counts()
    labels = df_counts.index.tolist()
    values = df_counts.values.tolist()
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
    plt.title('Test Results')
    plt.axis('equal')
    # Save the combined graph as bytes
    graph_combined1 = io.BytesIO()
    plt.savefig(graph_combined1)
    plt.close()

    return graph_combined1



def insulation_combined_graph(mf):
    mf = pd.read_csv("Insulate.csv")

    # Bar graph
    x = mf["Location"]
    y = mf["Nominal Circuit Voltage"]

    fig = plt.figure(figsize=(12, 6))  # Adjust the figsize as desired
    ax1 = fig.add_subplot(121)
    colors = ["#d9534f","#5bc0de","#5cb85c","#428bca"]                                       # Add more colors if needed
    ax1.bar(x, y, color=colors)
    ax1.set_xlabel("Location")
    ax1.set_ylabel("Nominal Circuit Voltage")
    ax1.set_title("Nominal Circuit Voltage by Location")
    

    # Pie chart
    earthing_system_counts = mf["Earthing System"].value_counts()
    ax2 = fig.add_subplot(122)
    colors = ["#5ac85a", "#dc0000"]
    ax2.pie(earthing_system_counts, labels=earthing_system_counts.index, autopct="%1.1f%%", colors=colors)
    ax2.set_title("Earthing System Distribution")
    ax2.axis("equal")
    

    graph_combined2 = io.BytesIO()
    plt.savefig(graph_combined2)
    plt.close()

    return graph_combined2



def phase_combined_graph(pf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = pf["Phase Sequence"]
    y = pf["V-L3-N"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Phase Sequence")
    plt.ylabel("V-L3-N")
    plt.title("Phase Sequence by V-L3-N")

    # Pie chart
    plt.subplot(122)
    pf['Result'] = phase_rang(pf)  # Ensure you have the phase_rang() function defined correctly
    pf_counts = pf['Result'].value_counts()
    labels = pf_counts.index.tolist()
    values = pf_counts.values.tolist()
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
    plt.axis('equal')
    plt.title('Test Results')
    graph_combined3 = io.BytesIO()
    plt.savefig(graph_combined3)
    plt.close()

    return graph_combined3


def polarity_combined_graph(af):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = af["Type of Supply"]
    y = af["Line to Neutral Voltage (V)"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Type of Supply")
    plt.ylabel("Line to Neutral Voltage (V)")
    plt.title("Type of Supply VS Line to Neutral Voltage (V)")

    # Pie chart
    plt.subplot(122)
    af['Result'] = polarity_rang(af.shape[0])  # Ensure you have the polarityrang() function defined correctly
    af_counts = af['Result'].value_counts()
    labels = af_counts.index.tolist()
    values = af_counts.values.tolist()
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
    plt.axis('equal')
    plt.title('Polarity Results')
    # Save the combined graph as bytes
    graph_combined4 = io.BytesIO()
    plt.savefig(graph_combined4)
    plt.close()

    return graph_combined4

def voltage_combined_graph(vf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = vf["Voltage Drop %"]
    y = vf["Calculated Voltage Drop (V)"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Voltage Drop %")
    plt.ylabel("Calculated Voltage Drop (V)")
    plt.title("Calculated Voltage Drop (V) VS Voltage Drop %")

    # Pie chart
    plt.subplot(122)
    vf['Result'] = voltage_rang(len(vf))  # Ensure you have the voltage_rang() function defined correctly
    vf_counts = vf['Result'].value_counts()
    labels = vf_counts.index.tolist()
    values = vf_counts.values.tolist()
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors)
    plt.axis('equal')
    plt.title('Test Results')
    graph_combined5 = io.BytesIO()
    plt.savefig(graph_combined5)
    plt.close()

    return graph_combined5


def residual_combined_graph(rf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    y= rf["Trip curve type"]
    x = rf["No. of Poles"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.ylabel("Trip curve type")
    plt.xlabel("No. of Poles")
    plt.title("Residual Current Device Test Results")

    # Pie chart
    plt.subplot(122)
    # rf['Result'] = residual_rang(rf.shape[0])  # Ensure you have the residual_rang() function defined correctly
    result_counts = rf["Result"].value_counts()
    labels = result_counts.index
    values = result_counts.values
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", shadow=False, startangle=90, colors=colors)
    plt.title("Residual Test Results")
    plt.axis("equal")  # Equal aspect ratio ensures that the pie is drawn as a circle
    graph_combined6 = io.BytesIO()
    plt.savefig(graph_combined6)
    plt.close()


    

    return graph_combined6



def Earth_combined_graph(ef):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    result_counts = ef["Result"].value_counts()
    colors = ["#d9534f","#5bc0de","#5cb85c","#428bca"]  # Add more colors if needed
    plt.bar(result_counts.index, result_counts.values, color=colors)  # Use 'color' instead of 'colors'
    plt.xlabel("Result")
    plt.ylabel("Count")
    plt.title("Earth Pit Electrode Test Results (Bar Graph)")

    # Pie chart
    plt.subplot(122)
    ef["Result"] = earthpit_rang(ef.shape[0])
    result_counts = ef["Result"].value_counts()
    labels = result_counts.index
    values = result_counts.values
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
    plt.title("Earth Pit Electrode Test Results (Pie Chart)")
    plt.axis('equal')
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined

def threephase_combined_graph(tf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = tf["ZeroSum Result"]
    y = tf["Zero Sum Current (mA)"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("ZeroSum Result")
    plt.ylabel("Zero Sum Current (mA)")
    plt.title("ZeroSum Result VS  Zero Sum Current (mA)")

    # Pie chart
    plt.subplot(122)
    result_counts = tf["ZeroSum Result"].value_counts()
    labels = result_counts.index
    values = result_counts.values
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
    plt.title("Three Phase Symmetry Test Results")
    plt.axis('equal')  # Equal aspect ratio ensures that the pie is drawn as a circle

    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined

def resc_combined_graph(jf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = jf["Conductor Type"]
    y = jf["Corrected Continuity Resistance (Ω)"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Conductor Type")
    plt.ylabel("Corrected Continuity Resistance")
    plt.title("Conductor Type VS Corrected Continuity Resistance")

    # Pie chart
    plt.subplot(122)
    jf['Result'] = resc_rang(jf)
    jf_counts = jf['Result'].value_counts()
    labels = jf_counts.index.tolist()
    values = jf_counts.values.tolist()
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
    plt.axis('equal')
    plt.title('Test Results')
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined



def func_ops_combined_graph(of):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = of["Interlock check"]
    y = of["SN"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Interlock check")
    plt.ylabel("SN")
    plt.title("SN VS Interlock check")

    # Pie chart
    plt.subplot(122)
    of['Result'] = func_ops_rang(of)
    of_counts = of['Result'].value_counts()
    labels = of_counts.index.tolist()
    values = of_counts.values.tolist()
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90,colors=colors)
    plt.axis('equal')
    plt.title('Test Results')
    # Save the combined graph as bytes
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined

def pat_combined_graph(bf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    y= bf["Earth Continuity (?)"]
    x = bf["Location"]
    colors = ["#d9534f", "#5bc0de", "#aa6f73", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.ylabel("Earth Continuity (?)")
    plt.xlabel("Location")
    plt.title("Location Location VS  Earth Continuity (?) ")

    # Pie chart
    plt.subplot(122)
    result_counts = bf["Overall Result"].value_counts()
    labels = result_counts.index
    values = result_counts.values
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", shadow=False, startangle=90, colors=colors)
    plt.title("Residual Test Results")
    plt.axis("equal")  # Equal aspect ratio ensures that the pie is drawn as a circle
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined

def socket_combined_graph(sf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x= sf["Facility Area"]
    y= sf["Upstream Breaker Rating (A)"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Facility Area")
    plt.ylabel("Upstream Breaker Rating (A)")
    plt.title("Facility Area VS  Upstream Breaker Rating (A)")

    # Pie chart
    plt.subplot(122)
    result_counts = sf2["Result"].value_counts()
    labels = result_counts.index
    values= result_counts.values

    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", shadow=False, startangle=90, colors=colors)
    plt.title("Residual Test Results")
    plt.axis("equal")  # Equal aspect ratio ensures that the pie is drawn as a circle
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined

def eli_test_combined_graph(gf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x= gf["Facility Area"]
    y= gf["Device Rating (A)"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Facility Area")
    plt.ylabel("Device Rating (A)")
    plt.title("Facility Area VS  Device Rating (A) ")

    # Pie chart
    plt.subplot(122)
    result_counts = gf2["Result"].value_counts()
    labels = result_counts.index
    values= result_counts.values

    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", shadow=False, startangle=90, colors=colors)
    plt.title("Residual Test Results")
    plt.axis("equal")  # Equal aspect ratio ensures that the pie is drawn as a circle
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined


def main():
    A = "floor.csv"
    df = pd.read_csv("floorfinal.csv")

    B = "Insulate.csv"
    mf = pd.read_csv("Insulate.csv")

    C = "phasesequence.csv"
    pf = pd.read_csv("phasesequence.csv")

    D = "pol.csv"
    af = pd.read_csv("pol.csv")

    E = "voltage.csv"
    vf = pd.read_csv("voltage_upd.csv")

    F = "residual.csv"
    rf = pd.read_csv("residual.csv")

    G = "earthpit.csv"
    ef = pd.read_csv(G)

    H="threephase.csv"
    tf=pd.read_csv("threephase.csv")
    I="threephasevalue.csv"
    tf2=pd.read_csv("threephasevalue.csv")

    J="resistance.csv"
    jf=pd.read_csv("resistance_updated.csv")

    H= "func_ops.csv"
    of = pd.read_csv(H)

    P="pat.csv"
    bf = pd.read_csv("pat.csv")

    Q="eli-socket.csv"
    sf = pd.read_csv("eli-socket.csv")
    R="sugg-max-eli.csv"
    fs = pd.read_csv("sugg-max-eli.csv")

    X="eli-test.csv"
    gf = pd.read_csv("eli-test.csv")
    Y="sugg-max-eli.csv"
    fg = pd.read_csv("sugg-max-eli.csv")



    doc = Document()
    
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(12)
    for section in doc.sections:
        section.left_margin = Inches(1)
    title = doc.add_heading("TESTING REPORT", 0)
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x6f, 0xa3, 0x15)

    section = doc.sections[0]
    header = section.header


    htable = header.add_table(1, 2, width=Inches(6))                                                  # Create a table with two cells for the pictures
    htable.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER                                                   # Configure the table properties
    htable.autofit = False

   
    cell1 = htable.cell(0, 0)                                                                        # Get the first cell in the table
    cell1.width = Inches(4)                                                                        # Adjust the width of the first cell

    left_header_image_path = "efficienergy-logo.jpg"                                                # Add the first picture to the first cell
    cell1_paragraph = cell1.paragraphs[0]
    cell1_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell1_run = cell1_paragraph.add_run()
    cell1_run.add_picture(left_header_image_path, width=Inches(1.5))

    # Get the second cell in the table
    cell2 = htable.cell(0, 1)
    cell2.width = Inches(3)                                                                     # Adjust the width of the second cell

    # Add the second picture to the second cell
    right_header_image_path = "secqr logo.png"                                              # Replace with the actual image file path
    cell2_paragraph = cell2.paragraphs[0]
    cell2_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    cell2_run = cell2_paragraph.add_run()
    cell2_run.add_picture(right_header_image_path, width=Inches(1.3))

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = "This Report is the Intellectual Property of M/s Efficienergi Consulting Pvt. Ltd. Plagiarism in Part or Full will be considered as theft of Intellectual property. The Information in this Report is to be treated as Confidential."
    for run in footer_paragraph.runs:
        run.font.name = "Calibri"                                                                                   # Replace with the desired font name
        run.font.size = Pt(7)                                                                                       # Replace with the desired font size

    doc.add_paragraph("FLOOR-RESISTANCE TEST")
    doc = resistance_table(df, doc)                                                                                     # Add a table of resistance data to the document
    graph_combined = flooresistance_combined_graph(df)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(3))                                                                                     # Add the resistance pie chart to the document

    doc.add_paragraph("INSULATION TEST")
    doc = insulation_table(mf, doc)                                                                                           # Add a table of insulation data to the document
    graph_combined = insulation_combined_graph(mf)
    doc.add_picture(graph_combined,width=Inches(8), height=Inches(4))                                                         # Add the insulation pie chart to the document

    doc.add_paragraph("PHASE SEQUENCE TEST")
    doc = phase_table(pf, doc)                                                                                                   # Add a table of phase sequence data to the document
    graph_combined = phase_combined_graph(pf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))                                                            # Add the phase sequence pie chart to the document

    doc.add_paragraph("POLARITY TEST")
    doc = polarity_table(af, doc)                                                                                                        # Add a table of polarity data to the document
    graph_combined = polarity_combined_graph(af)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))                                                                   # Add the polarity pie chart to the document

    doc.add_paragraph("VOLTAGE DROP TEST")
    doc = voltage_table(vf, doc)                                                                                                   # Add a table of voltage drop data to the document
    graph_combined = voltage_combined_graph(vf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))                                                                # Add the voltage drop pie chart to the document

    doc.add_paragraph("Residual Current Device Test")
    doc = residual_table(rf, doc)                                                                                                     # Add a table of residual current device data to the document
    graph_combined = residual_combined_graph(rf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))                                                                                                            # Generate a graph of residual current device data

    doc.add_paragraph("EARTH PIT  RESISTANCE TEST")
    doc = earthpit_table(ef, doc)                                                                                                        # Add a table of earth pit resistance data to the document
    graph_combined1 = Earth_combined_graph(ef)
    doc.add_picture(graph_combined1, width=Inches(8), height=Inches(4))


    doc.add_paragraph("THREE PHASE SYMMETRY TEST")
    doc = threephase_table(tf, doc)                                                                                                      # Add a table of three-phase symmetry data to the document
    graph = threephase_combined_graph(tf)                                                                                                         # Generate a graph of three-phase symmetry data
    graph_combined = threephase_combined_graph(tf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))                                                                                                              # Add the three-phase symmetry pie chart to the document

    doc.add_paragraph("RESISTANCE CONDUCTOR TEST")
    doc = resc_table(jf, doc)                                                                                                      # Add a table of three-phase symmetry data to the document
    graph_combined = resc_combined_graph(jf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    doc.add_heading('FUNCTIONS AND OPERATION TEST')
    doc = func_ops_table(of, doc)
    graph_combined = func_ops_combined_graph(of)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))


    doc.add_heading('PAT TEST')
    doc = pat_table(bf, doc)
    graph_combined = pat_combined_graph(bf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4)) 

    doc.add_paragraph("ELI SOCKET TEST")  
    doc = create_eli_table1(sf1, doc)
    doc.add_paragraph("\n")
    doc = create_eli_table2(sf2, doc)                                                                                                 # Add a table of voltage drop data to the document
    graph_combined = socket_combined_graph(sf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))   

    
    doc.add_paragraph("ELI CIRCUIT BREAKER TEST") 
    doc = eli_test_table1(gf1, doc)
    doc.add_paragraph("\n")
    doc = eli_test_table2(gf2, doc)
    doc.add_paragraph("ELI  TEST")                                                                                                   # Add a table of voltage drop data to the document
    graph_combined = eli_test_combined_graph(gf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))
    doc.save("ELI_Report.docx")   




    doc.save("scriptreport.docx")                                                                                                        # Save the Word document with all the added content

   

main()