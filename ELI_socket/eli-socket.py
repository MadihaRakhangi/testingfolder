import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
import io
import numpy as np
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tabulate import tabulate
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

sf = pd.read_csv("eli-socket.csv")
fs = pd.read_csv("sugg-max-eli.csv")


sf1 = sf[
    [
        "SN",
        "Socket Name",
        "Location",
        "Facility Area",
        "Socket Type",
        "Earthing Configuration",
        "Upstream Breaker Rating (A)",
        "Upstream Breaker Make",
        "Upstream Breaker Type",
        "Trip Curve",
    ]
]

sf2 = sf[
    [
        "SN",
        "Socket Name",
        "Socket Rating (A)",
        "Socket Type",
        "No. of Phases",
        "V_LN",
        "V_LE",
        "V_NE",
        "L1-ELI",
        "L2-ELI",
        "L3-ELI",
        "Psc (kA)",
    ]
]

sf_filled = sf.fillna("")
sf["Upstream Breaker Rating (A)"] = sf["Upstream Breaker Rating (A)"].astype(int)

Device_Rating = sf["Upstream Breaker Rating (A)"]
No_phase = sf["No. of Phases"]
T_Curve = sf["Trip Curve"]
new_column = []
result_column = []
P = 0
K = 0
TMS = 1
TDS = 1


def socket_result1(sf1, sf2):
    I = Is * ((((K * TMS) / Td) + 1) ** (1 / P))
    if row["Earthing Configuration"] == "TN":
        IEC_val_TN = row["V_LE"] / I
        new_column.append(round(IEC_val_TN, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= IEC_val_TN
                and row["L2-ELI"] <= IEC_val_TN
                and row["L3-ELI"] <= IEC_val_TN
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= IEC_val_TN:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")
    elif row["Earthing Configuration"] == "TT":
        IEC_val_TT = 50 / I
        new_column.append(round(IEC_val_TT, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= IEC_val_TT
                and row["L2-ELI"] <= IEC_val_TT
                and row["L3-ELI"] <= IEC_val_TT
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= IEC_val_TT:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")


def socket_result2(sf1, df2):
    I = Is * (((((A / ((Td / TDS) - B)) + 1)) ** (1 / p)))
    if row["Earthing Configuration"] == "TN":
        IEEE_val_TN = row["V_LE"] / I
        new_column.append(round(IEEE_val_TN, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= IEEE_val_TN
                and row["L2-ELI"] <= IEEE_val_TN
                and row["L3-ELI"] <= IEEE_val_TN
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= IEEE_val_TN:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")
    elif row["Earthing Configuration"] == "TT":
        IEEE_val_TT = 50 / I
        new_column.append(round(IEEE_val_TT, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= IEEE_val_TT
                and row["L2-ELI"] <= IEEE_val_TT
                and row["L3-ELI"] <= IEEE_val_TT
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= IEEE_val_TT:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")


for index, row in sf.iterrows():
    if row["Upstream Breaker Type"] == "MCB":
        rating = row[8]
        trip = row[13]  # Assuming the "Trip Curve" column is at index 10
        result_row = fs[fs["Device Rating (A)"] == rating]
        if trip in result_row.columns:
            val_MCB = result_row[trip].values[0]
        else:
            val_MCB = (
                0  # Set a default value when 'Trip Curve' value is not found in sugg-max-eli.csv
            )
        new_column.append(round(val_MCB, 2))

        if row["No. of Phases"] == 3:
            if row["L1-ELI"] <= val_MCB and row["L2-ELI"] <= val_MCB and row["L3-ELI"] <= val_MCB:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= val_MCB:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")

    elif (
        row["Upstream Breaker Type"] in ["RCD", "RCBO", "RCCB"]
        and row["Earthing Configuration"] == "TN"
    ):
        rccb_val_TN = (row["V_LE"] / row["Upstream Breaker Sensitivity"]) * 1000
        new_column.append(round(rccb_val_TN, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= rccb_val_TN
                and row["L2-ELI"] <= rccb_val_TN
                and row["L3-ELI"] <= rccb_val_TN
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= rccb_val_TN:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")

    elif (
        row["Upstream Breaker Type"] in ["RCD", "RCBO", "RCCB"]
        and row["Earthing Configuration"] == "TT"
    ):
        rccb_val_TT = (50 / row["Upstream Breaker Sensitivity"]) * 1000
        new_column.append(round(rccb_val_TT, 4))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI"] <= rccb_val_TT
                and row["L2-ELI"] <= rccb_val_TT
                and row["L3-ELI"] <= rccb_val_TT
            ):
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI"] <= rccb_val_TT:
                result_column.append("Pass")
            else:
                result_column.append("Fail")
        else:
            result_column.append("N/A")

    elif row["Upstream Breaker Type"] == "MCCB":
        if row["Type of Circuit Location"] == "Final":
            Td = 0.4
        elif row["Type of Circuit Location"] == "Distribution":
            Td = 5
        Is = row["Upstream Breaker Rating (A)"]
        if row[13] == "IEC Standard Inverse":
            P = 0.02
            K = 0.14
            socket_result1(sf1, sf2)
        elif row[13] == "IEC Very Inverse":
            P = 1
            K = 13.5
            socket_result1(sf1, sf2)
        elif row[13] == "IEC Long-Time Inverse":
            P = 1
            K = 120
            socket_result1(sf1, sf2)
        elif row[13] == "IEC Extremely Inverse":
            P = 2
            K = 80
            socket_result1(sf1, sf2)
        elif row[13] == "IEC Ultra Inverse":
            P = 2.5
            K = 315.2
            socket_result1(sf1, sf2)
        elif row[13] == "IEEE Moderately Inverse":
            A = 0.0515
            B = 0.114
            p = 0.02
            socket_result1(sf1, sf2)
        elif row[13] == "IEEE Very Inverse":
            A = 19.61
            B = 0.491
            p = 2
            socket_result1(sf1, sf2)
        elif row[13] == "IEEE Extremely Inverse":
            A = 28.2
            B = 0.1217
            p = 2
            socket_result2(sf1, sf2)


new_column = pd.Series(new_column[: len(sf2)], name="Suggested Max ELI (Ω)")
sf2.loc[:, "Suggested Max ELI (Ω)"] = new_column
sf2.loc[:, "Suggested Max ELI (Ω)"] = sf2["Suggested Max ELI (Ω)"].apply(lambda x: "{:.2f}".format(x))
result_column = pd.Series(result_column[: len(sf2)], name="Result")
sf2["Result"] = result_column


def create_eli_table1(sf1, doc):
    sf1 = sf1.fillna("")
    doc.add_heading("Earth Loop Impedance Test - Circuit Breaker", level=1)
    table_data = sf1.iloc[:, 0:]
    table_str = tabulate(table_data, headers="keys", tablefmt="pipe")
    num_rows, num_cols = table_data.shape[0], table_data.shape[1]
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.2,
        1: 0.5,
        2: 0.75,
        3: 0.6,
        4: 0.8,
        5: 0.8,
        6: 0.6,
        7: 0.8,
        8: 1,
        9: 1.25,
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths.get(j, 1))
        table.cell(0, j).paragraphs[
            0
        ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.paragraphs[
                0
            ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center
            if j == num_cols - 1:  # Apply background color only to the Result column
                result_cell = cell
                if value == "Pass":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls("w"))
                    )  # Green color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
                elif value == "Fail":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls("w"))
                    )  # Red color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"

    return doc


def create_eli_table2(sf2, doc):
    table_data = sf2.iloc[:, 0:]
    table_str = tabulate(table_data, headers="keys", tablefmt="pipe")
    num_rows, num_cols = table_data.shape[0], table_data.shape[1]
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.2,
        1: 0.44,
        2: 0.55,
        3: 0.59,
        4: 0.6,
        5: 0.58,
        6: 0.45,
        7: 0.47,
        8: 0.45,
        9: 0.41,
        10: 0.41,
        11: 0.41,
        12: 0.62,
        13: 0.8,
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths.get(j, 1))
        table.cell(0, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        table.cell(0, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if j == num_cols - 1:  # Apply background color only to the Result column
                result_cell = cell
                if value == "Pass":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls("w"))
                    )  # Green color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
                elif value == "Fail":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls("w"))
                    )  # Red color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"

    return doc

def socket_combined_graph(sf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x= sf["Facility Area"]
    y= sf["Upstream Breaker Rating (A)"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.xlabel("Facility Area")
    plt.ylabel("Upstream Breaker Rating (A)")
    plt.title("Facility Area VS  Upstream Breaker Rating (A)")

    # Pie chart
    plt.subplot(122)
    result_counts = sf2["Result"].value_counts()
    labels = result_counts.index
    values= result_counts.values

    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", shadow=False, startangle=90, colors=colors)
    plt.title("Residual Test Results")
    plt.axis("equal")  # Equal aspect ratio ensures that the pie is drawn as a circle
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined

def main():
    doc = Document()
    doc = create_eli_table1(sf1, doc)
    doc.add_paragraph("\n")
    doc = create_eli_table2(sf2, doc)
    doc.add_paragraph("ELI SOCKET TEST")                                                                                                   # Add a table of voltage drop data to the document
    graph_combined = socket_combined_graph(sf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))      
    doc.save("ELI_Socket.docx")

main()




