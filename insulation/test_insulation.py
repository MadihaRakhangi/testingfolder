import pytest
from Insulationtest import insualtion_result , insulationrang , insulation_graph , insulation_pie , insulation_table
import pandas as pd
import io
import os
import matplotlib as plt
from matplotlib.image import imread
from matplotlib.pyplot import imshow

from matplotlib.testing.decorators import image_comparison
import tkinter
import filecmp
from docx import Document


test_df=pd.read_csv("Insulate.csv")

@pytest.fixture
def sample_dataframe():
                                                           
    df = pd.DataFrame({                                           # Define a sample dataframe for testing
        "Nominal Circuit Voltage": [10, 100, 400],
        "Test Voltage": [250, 500, 1900],
        "Insulator Resistance": [0.5, 1.8, 0.9]
    })
    return df



def test_result():
    assert insualtion_result(10, 250, 0.5) == "Satisfactory"
    assert insualtion_result(100, 500, 0.8) == "Unsatisfactory"


def test_rang():
    df = pd.DataFrame({                                  #create  a test dataframe
         "Nominal Circuit Voltage": [10, 100, 400],
         "Test Voltage": [250, 500, 1900],
         "Insulator Resistance": [0.5, 1.8, 0.9]
    })
    assert insulationrang(df.shape[0]) == ["Satisfactory", "Satisfactory", "Unsatisfactory"]
    

def test_create_table():
    df = pd.DataFrame({                               #create  a test dataframe
        "Serial number": [1, 2, 3],
        "Location": ["A", "B", "C"],
        "Nominal Circuit Voltage": [100, 200, 300]
    })

    doc = Document()                                  # Create a Document object
    doc = insulation_table(df, doc)                      
    temp_file = "temp_doc.docx"                        # Save the document as a temporary file
    temp_file = "temp_doc.docx"
    doc.save(temp_file)
    doc_read = Document(temp_file)                     # Read the temporary file and check if the table exists
    tables = doc_read.tables
    assert len(tables) == 1                             # Check if the first table in the document matches the DataFrame shape
    table = tables[0]
    assert len(table.rows) == df.shape[0] + 1            # Include header row
    assert len(table.columns) == df.shape[1]+1           #include the result coloumn
    os.remove(temp_file) 


def test_graph(sample_dataframe):
    graph2 = insulation_graph(sample_dataframe)
    # Add assertions to check the generated graph

def test_pie(sample_dataframe):
    graph4 = insulation_pie(sample_dataframe)
    # Add assertions to check the generated graph


