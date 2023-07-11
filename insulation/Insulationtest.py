import pandas as pd
import matplotlib.pyplot as plt
import docx
import csv
import io
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

mf = pd.read_csv("Insulate.csv")


def insualtion_result(Nom_CVolt, T_Volt, Insu_R):
    if Nom_CVolt <= 50:
        if Insu_R >= 0.5 and T_Volt == 250:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    elif 50 < Nom_CVolt <= 500:
        if Insu_R >= 1 and T_Volt == 500:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    elif Nom_CVolt > 500:
        if Insu_R >= 1 and T_Volt == 1000:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    else:
        return "Invalid input"
    

    


def insulationrang(length):
    res2 = []
    for row in range(length):  # Adjusted the range to start from 0
        Nom_CVolt = mf.iloc[row, 7]
        T_Volt = mf.iloc[row, 9]
        Insu_R = mf.iloc[row, 13]
        res2.append(insualtion_result(Nom_CVolt, T_Volt, Insu_R))
        print(Nom_CVolt)
    return res2

def insulation_table(mf, doc):
    table_data = mf.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)  # Add +1 for the "Result" column
    table.style = "Table Grid"
    table.autofit = False
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.55,
        3: 0.54,
        4: 0.38,
        5: 0.4,
        6: 0.5,
        7: 0.48,
        8: 0.5,
        9: 0.43,
        10: 0.4,
        11: 0.4,
        12: 0.4,
        13: 0.4,
        num_cols: 0.8  # Width for the "Result" column
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)
    Results = insulationrang(num_rows)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(column_widths[num_cols])  # Set width for the "Result" column
    for i in range(0, num_rows):
        res_index = i - 1
        cell = table.cell(i + 1, num_cols)
        cell.text = Results[res_index]
        
        if Results[res_index] == "Satisfactory":
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="#5ac85a"/>'.format(nsdecls("w"))
            )  # Green color
            cell._tc.get_or_add_tcPr().append(shading_elm)
        elif Results[res_index] == "Unsatisfactory":
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="#dc0000"/>'.format(nsdecls("w"))
            )  # Red color
            cell._tc.get_or_add_tcPr().append(shading_elm)
    
    font_size = 6.5

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    return doc

def insulation_combined_graph(mf):
    mf = pd.read_csv("Insulate.csv")

    # Bar graph
    x = mf["Location"]
    y = mf["Nominal Circuit Voltage"]

    fig = plt.figure(figsize=(12, 6))  # Adjust the figsize as desired
    ax1 = fig.add_subplot(121)
    colors = ["#d9534f","#5bc0de","#5cb85c","#428bca"]                                       # Add more colors if needed
    ax1.bar(x, y, color=colors)
    ax1.set_xlabel("Location")
    ax1.set_ylabel("Nominal Circuit Voltage")
    ax1.set_title("Nominal Circuit Voltage by Location")
    

    # Pie chart
    earthing_system_counts = mf["Earthing System"].value_counts()
    ax2 = fig.add_subplot(122)
    colors = ["#5ac85a", "#dc0000"]
    ax2.pie(earthing_system_counts, labels=earthing_system_counts.index, autopct="%1.1f%%", colors=colors)
    ax2.set_title("Earthing System Distribution")
    ax2.axis("equal")
    

    graph_combined2 = io.BytesIO()
    plt.savefig(graph_combined2)
    plt.close()

    return graph_combined2




def main():
    mf = pd.read_csv("Insulate.csv")
    doc = docx.Document()
    doc = insulation_table(mf, doc)
    
    graph_combined = insulation_combined_graph(mf)
    doc.add_picture(graph_combined,width=Inches(8), height=Inches(4)) 
    doc.save("outputTEST.docx")


main()