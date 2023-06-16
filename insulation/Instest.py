import pandas as pd
import matplotlib.pyplot as plt
import docx
import csv
import io
from docx.shared import Inches
from docx.shared import Pt

df = pd.read_csv("Insulate.csv")

Sn = df.iloc[:, [0]]  # Serial number
Loc = df.iloc[:, [1]]  # Location
P_Loc = df.iloc[:, [2]]  # Parent Location
Up_DName = df.iloc[:, [3]]  # Upstream Device Name
N_Poles = df.iloc[:, [4]]  # Number of Poles
SPD_A = df.iloc[:, [5]]  # SPD Applicable
E_Sys = df.iloc[:, [6]]  # Earthing System
Nom_CVolt = df.iloc[:, [7]]  # Nominal Circuit Voltage
Mes_T = df.iloc[:, [8]]  # Measurement Terminal
T_Volt = df.iloc[:, [9]]  # Test Voltage
Cond_T = df.iloc[:, [10]]  # Conductor Type
Cond_S = df.iloc[:, [11]]  # Conductor Size
Insu_T = df.iloc[:, [12]]  # Insulator Type
Insu_R = df.iloc[:, [13]]  # Insulator Resistance


def result(Nom_CVolt, T_Volt, Insu_R):
    if Nom_CVolt <= 50:
        if Insu_R >= 0.5 and T_Volt == 250:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    elif 50 < Nom_CVolt <= 500:
        if Insu_R >= 1 and T_Volt == 500:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    elif Nom_CVolt > 500:
        if Insu_R >= 1 and T_Volt == 1000:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    else:
        return "Invalid input"

def rang(length):
    res = []
    for row in range(length):  # Adjusted the range to start from 0
        Nom_CVolt = df.iloc[row, 7]
        T_Volt = df.iloc[row, 9]
        Insu_R = df.iloc[row, 13]
        res.append(result(Nom_CVolt, T_Volt, Insu_R))
        print(Nom_CVolt)
    return res

# def rang(length):
#     res = []
#     for row in range(1, length):
#         Nom_CVolt = df.iloc[row, 7]
#         T_Volt = df.iloc[row, 9]
#         Insu_R = df.iloc[row, 13]
#         res.append(result(Nom_CVolt, T_Volt, Insu_R))
#         print(Nom_CVolt)
#     return res


# def create_table(df, doc):
#     table_data = df.iloc[:, 0:]
#     num_rows, num_cols = table_data.shape
#     table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
#     table.style = "Table Grid"
#     table.autofit = False
#     column_widths = {
#         0: 0.2,
#         1: 0.51,
#         2: 0.55,
#         3: 0.54,
#         4: 0.38,
#         5: 0.56,
#         6: 0.5,
#         7: 0.48,
#         8: 0.71,
#         9: 0.43,
#         10: 0.56,
#         11: 0.56,
#         12: 0.56,
#         13: 0.6,
#         
#     }
#     for j, col in enumerate(table_data.columns):
#         table.cell(0, j).text = col
#         table.cell(0, j).width = Inches(column_widths[j])
#     for i, row in enumerate(table_data.itertuples(), start=1):
#         for j, value in enumerate(row[1:], start=0):
#             table.cell(i, j).text = str(value)
#     Results = rang(num_rows)
#     table.cell(0, num_cols).text = "Result"
#     table.cell(0, num_cols).width = Inches(0.8)
#     for i in range(0, num_rows):
#         res_index = i - 1
#         table.cell(i + 1, num_cols).text = Results[res_index]
#     font_size = 6.5

#     for row in table.rows:
#         for cell in row.cells:
#             for paragraph in cell.paragraphs:
#                 for run in paragraph.runs:
#                     run.font.size = Pt(font_size)

#     return doc

def create_table(df, doc):
    table_data = df.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)  # Add +1 for the "Result" column
    table.style = "Table Grid"
    table.autofit = False
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.55,
        3: 0.54,
        4: 0.38,
        5: 0.4,
        6: 0.5,
        7: 0.48,
        8: 0.5,
        9: 0.43,
        10: 0.4,
        11: 0.4,
        12: 0.4,
        13: 0.4,
        num_cols: 0.8  # Width for the "Result" column
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)
    Results = rang(num_rows)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(column_widths[num_cols])  # Set width for the "Result" column
    for i in range(0, num_rows):
        res_index = i - 1
        table.cell(i + 1, num_cols).text = Results[res_index]
    font_size = 6.5

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    return doc

def graph(df):
    df = pd.read_csv("Insulate.csv")
    x = df["Location"]
    y = df["Nominal Circuit Voltage"]
    plt.bar(x, y)
    plt.xlabel("Location")
    plt.ylabel("Nominal Circuit Voltage")
    plt.title("Nominal Circuit Voltage by Location")
    plt.savefig("chart.png")
    graph = io.BytesIO()
    plt.savefig(graph)
    plt.close()
    return graph


def graph_pie(df):
    df = pd.read_csv("Insulate.csv")
    earthing_system_counts = df["Earthing System"].value_counts()
    plt.pie(earthing_system_counts, labels=earthing_system_counts.index, autopct="%1.1f%%")
    plt.axis("equal")
    plt.title("Earthing System Distribution")
    plt.savefig("pie_chart.png")
    graph_pie = io.BytesIO()
    plt.savefig(graph_pie)
    plt.close()
    return graph_pie


def create_scatter_plot(df, doc):
    fig, ax = plt.subplots()
    ax.scatter(df["Insulation Resistance (MΩ)"], df["Test Voltage (V)"])
    ax.set_xlabel("Insulation Resistance (MΩ)")
    ax.set_ylabel("Test Voltage (V)")
    ax.set_title("Insulation Resistance vs Test Voltage")

    # Save the plot to a BytesIO object
    image_stream = io.BytesIO()
    plt.savefig(image_stream, format="png")
    plt.close(fig)

    # Add the plot image to the Word document
    doc.add_picture(image_stream)
    return doc


def main():
    df = pd.read_csv("Insulate.csv")
    doc = docx.Document()
    doc = create_table(df, doc)

    histo_graph = graph(df)
    doc.add_picture(histo_graph)

    pie_graph = graph_pie(df)
    doc.add_picture(pie_graph)
    doc = create_scatter_plot(df, doc)
    doc.save("outputTEST.docx")


main()