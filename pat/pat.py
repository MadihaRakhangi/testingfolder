import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
import pandas as pd
import io
import matplotlib.pyplot as plt
from docx import Document
import numpy as np
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tabulate import tabulate
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

P="pat.csv"
bf = pd.read_csv("pat.csv")



def pat_table(bf, doc):
    table_data = bf.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.2,
        1: 0.5,
        2: 0.6,
        3: 0.6,
        4: 0.8,
        5: 0.7,
        6: 0.6,
        7: 0.6,
        8: 0.6,
        9: 0.7,
        10: 0.48,
        11: 0.6,
        12: 0.9,
        13: 0.9,
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths.get(j, 1))
        table.cell(0, j).paragraphs[
            0
        ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.paragraphs[
                0
            ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center
            if j == num_cols - 1:  # Apply background color only to the Result column
                result_cell = cell
                if value == "PASS":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="#5ac85a"/>'.format(nsdecls("w"))
                    )  # Green color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
                elif value == "FAIL":
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="#dc0000"/>'.format(nsdecls("w"))
                    )  # Red color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"

    return doc

def pat_combined_graph(bf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    y= bf["Earth Continuity (?)"]
    x = bf["Location"]
    colors = ["#d9534f", "#5bc0de", "#aa6f73", "#428bca"]
    plt.bar(x, y, color=colors)
    plt.ylabel("Earth Continuity (?)")
    plt.xlabel("Location")
    plt.title("Location Location VS  Earth Continuity (?) ")

    # Pie chart
    plt.subplot(122)
    result_counts = bf["Overall Result"].value_counts()
    labels = result_counts.index
    values = result_counts.values
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", shadow=False, startangle=90, colors=colors)
    plt.title("Residual Test Results")
    plt.axis("equal")  # Equal aspect ratio ensures that the pie is drawn as a circle
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined


def main():
    doc=Document()
    doc = pat_table(bf, doc)
    graph_combined = pat_combined_graph(bf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4)) 
    doc.save("pat_Report.docx")

main()
