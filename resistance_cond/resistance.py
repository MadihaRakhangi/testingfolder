import pandas as pd
import matplotlib.pyplot as plt
import docx
from docx import Document
import io
from docx.shared import Inches
from docx.shared import Pt

E="resistance.csv"
jf=pd.read_csv(E)

def alpha(Cond_T):                               # Function to calculate alpha value based on conductor typ      
    if Cond_T == "Al":
        return 0.0038
    elif Cond_T == "Cu":
        return 0.00429
    elif Cond_T == "GI":
        return 0.00641
    elif Cond_T == "SS":
        return 0.003
    else:
        return None

# Calculate Corrected Continuity Resistance
jf['Corrected Continuity Resistance (Ω)'] = jf['Continuity Resistance (?)'] - jf['Lead Internal Resistance (?)']

# Calculate Specific Conductor Resistance at 30°C
alpha_values = jf['Conductor Type'].apply(alpha)
jf['Specific Conductor Resistance (MΩ/m) at 30°C'] = jf['Corrected Continuity Resistance (Ω)'] / (
    1 + alpha_values * (jf['Conductor Temperature (°C)'] - 30))
jf['Specific Conductor Resistance (MΩ/m) at 30°C'] /= 1000000
jf['Specific Conductor Resistance (MΩ/m) at 30°C'] = jf['Specific Conductor Resistance (MΩ/m) at 30°C'].apply(
    lambda x: format(x, 'E'))
jf.to_csv('resistance_updated.csv', index=False)


def resc_result(Conti, C_ContR):
    if C_ContR <= 1:
        if Conti == "Yes":
            return "Pass"
        elif Conti == "No":
            return "Check Again"
        else:
            return "Invalid"
    elif C_ContR > 1:
        if Conti == "Yes":
            return "Fail"
        else:
            return "Fail"


def resc_rang(jf):
    res5 = []
    for row in range(0, len(jf)):
        Conti = jf.iloc[row]['Is Continuity found?']
        C_ContR = jf.iloc[row]['Corrected Continuity Resistance (Ω)']
        res5.append(resc_result(Conti, C_ContR))
    return res5


def resc_table(jf, doc):
    table_data = jf.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False
    column_widths = {
        0: 0.2,
        1: 0.51,
        2: 0.55,
        3: 0.54,
        4: 0.38,
        5: 0.56,
        6: 0.5,
        7: 0.48,
        8: 0.71,
        9: 0.43,
        10: 0.56,
        11: 0.56, 
        12:0.5   # Add this line
    }
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)
    Results = resc_rang(jf)
    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(0.8)
    for i in range(num_rows):
        res_index = i
        table.cell(i + 1, num_cols).text = Results[res_index]
    font_size = 6.5

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    return doc


def resc_combined_graph(jf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = jf["Conductor Type"]
    y = jf["Corrected Continuity Resistance (Ω)"]
    plt.bar(x, y)
    plt.xlabel("Conductor Type")
    plt.ylabel("Corrected Continuity Resistance")
    plt.title("Conductor Type VS Corrected Continuity Resistance")

    # Pie chart
    plt.subplot(122)
    jf['Result'] = resc_rang(jf)
    jf_counts = jf['Result'].value_counts()
    labels = jf_counts.index.tolist()
    values = jf_counts.values.tolist()
    colors = ["#5ac85a", "#dc0000"]
    plt.pie(values, labels=labels, autopct='%1.1f%%', startangle=90,colors=colors)
    plt.axis('equal')
    plt.title('Test Results')
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined



def main():
    doc = Document()
    doc.add_heading('RESISTANCE CONDUCTOR TEST', 0)
    doc = resc_table(jf, doc)
    graph_combined = resc_combined_graph(jf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))
    doc.save("resfinal.docx")


main()
