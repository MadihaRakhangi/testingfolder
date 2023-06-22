import pytest
import pandas as pd
import numpy as np
from docx import Document
import os
from PIL import Image
from resistance import alpha, resc_result, resc_rang, resc_table, resc_graph, resc_pie


@pytest.fixture
def sample_dataframe():
    df = pd.DataFrame({
        'Is Continuity found?': ["Yes", "No"],
        'Conductor Type': ['Al', 'Cu'],
        'Conductor Temperature (°C)': [20, 40],
        'Continuity Resistance (Ω)': [1.1, 2.5],
        'Corrected Continuity Resistance (Ω)': [1, 2.3],
        'Lead Internal Resistance (Ω)': [0.05, 0.1],
    })
    return df


def test_alpha():
    assert alpha('Al') == 0.0038
    assert alpha('Cu') == 0.00429
    assert alpha('GI') == 0.00641
    assert alpha('SS') == 0.003
    assert alpha('Invalid') is None


def test_resc_result():
    assert resc_result('Yes', 0.9) == 'Pass'
    assert resc_result('No', 0.9) == 'Check Again'
    assert resc_result('Yes', 1.2) == 'Fail'
    assert resc_result('Invalid', 0.9) == 'Invalid'


def test_resc_rang(sample_dataframe):
    df = pd.DataFrame({
        "Is Continuity found?": ["Yes", "No"],
        "Corrected Continuity Resistance (Ω)": [0.9, 1.2],
    })
    assert resc_rang(df) == ["Pass", "Fail"]
    


def test_resc_table(sample_dataframe):
    doc = Document()
    doc = resc_table(sample_dataframe, doc)
    temp_file = "temp_doc.docx"
    doc.save(temp_file)
    doc_read = Document(temp_file)
    tables = doc_read.tables
    assert len(tables) == 1
    table = tables[0]
    assert len(table.rows) == sample_dataframe.shape[0] + 1
    assert len(table.columns) == sample_dataframe.shape[1] + 1
    os.remove(temp_file)


def test_resc_graph(sample_dataframe):
    graph8 = resc_graph(sample_dataframe)
    image = Image.open(graph8)
    image_array = np.array(image)
    assert image_array.shape[0] > 0
    assert image_array.shape[1] > 0
    graph8.close()


def test_resc_pie(sample_dataframe):
    graph = resc_pie(sample_dataframe)
    image = Image.open(graph)
    image_array = np.array(image)
    assert image_array.shape[0] > 0
    assert image_array.shape[1] > 0
    graph.close()
